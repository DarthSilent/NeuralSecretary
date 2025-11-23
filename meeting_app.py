import os
import sys
import logging
from logging.handlers import RotatingFileHandler

# --- 1. НАСТРОЙКА ПУТЕЙ ДЛЯ PORTABLE / PYINSTALLER ---
if getattr(sys, "frozen", False):
    internal_path = sys._MEIPASS
    # Если это EXE, храним данные в Документах пользователя, чтобы не требовать прав админа
    external_path = os.path.join(os.path.expanduser("~"), "Documents", "MeetingApp")
    os.makedirs(external_path, exist_ok=True)
else:
    internal_path = os.getcwd()
    external_path = os.getcwd()

os.environ["PATH"] += os.pathsep + internal_path
os.chdir(external_path)

# Add nvidia CUDA 12 DLL directories to PATH for faster-whisper
try:
    import site
    site_packages = site.getsitepackages()
    for sp in site_packages:
        nvidia_cublas_bin = os.path.join(sp, "nvidia", "cublas", "bin")
        nvidia_cudnn_bin = os.path.join(sp, "nvidia", "cudnn", "bin")
        if os.path.exists(nvidia_cublas_bin):
            os.environ["PATH"] = nvidia_cublas_bin + os.pathsep + os.environ["PATH"]
        if os.path.exists(nvidia_cudnn_bin):
            os.environ["PATH"] = nvidia_cudnn_bin + os.pathsep + os.environ["PATH"]
except Exception:
    pass

import warnings
warnings.filterwarnings("ignore")

import queue
import json
import requests
import re
import zipfile
import shutil
import subprocess
import uuid
from io import BytesIO
from datetime import datetime, timedelta
import threading
import time

# --- 2. НАСТРОЙКА ЛОГИРОВАНИЯ ---
LOG_FILE = "meeting_app.log"
logger = logging.getLogger("MeetingApp")
logger.setLevel(logging.DEBUG)

file_handler = RotatingFileHandler(LOG_FILE, maxBytes=5*1024*1024, backupCount=3, encoding='utf-8')
file_handler.setLevel(logging.DEBUG)
file_formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
file_handler.setFormatter(file_formatter)
logger.addHandler(file_handler)

console_handler = logging.StreamHandler()
console_handler.setLevel(logging.INFO)
console_formatter = logging.Formatter('%(levelname)s: %(message)s')
console_handler.setFormatter(console_formatter)
logger.addHandler(console_handler)

logger.info("========== Meeting App Started ==========")

import tkinter as tk
from tkinter import filedialog, messagebox

import customtkinter as ctk
import sounddevice as sd
import soundfile as sf
import numpy as np

import torch
import torchaudio
from dataclasses import dataclass

# --- MONKEYPATCH: FIX PYANNOTE.AUDIO MISSING TORCHCODEC ---
# Pyannote 3.3+ requires torchcodec, which is hard to install on Windows.
# We inject a compatibility layer using torchaudio.
try:
    import pyannote.audio.core.io
    
    # Always apply patch if AudioDecoder is missing OR if we want to force our logging version
    # But to be safe, we check if it's missing or if it's our own class
    if not hasattr(pyannote.audio.core.io, "AudioDecoder") or getattr(pyannote.audio.core.io.AudioDecoder, "__module__", "") == __name__:
        logger.info("Applying AudioDecoder monkeypatch for pyannote.audio...")
        
        @dataclass
        class AudioStreamMetadata:
            sample_rate: int
            duration_seconds_from_header: float
            num_channels: int

        @dataclass
        class AudioSamples:
            data: torch.Tensor
            sample_rate: int

        class AudioDecoder:
            def __init__(self, path):
                logger.debug(f"AudioDecoder init for {path}")
                self.path = path
                try:
                    self.info = torchaudio.info(path)
                    logger.debug(f"torchaudio.info success: {self.info}")
                except Exception as e:
                    logger.error(f"Error getting info for {path}: {e}")
                    raise

            @property
            def metadata(self):
                return AudioStreamMetadata(
                    sample_rate=self.info.sample_rate,
                    duration_seconds_from_header=self.info.num_frames / self.info.sample_rate,
                    num_channels=self.info.num_channels
                )

            def get_all_samples(self):
                logger.debug(f"get_all_samples for {self.path}")
                try:
                    waveform, sr = torchaudio.load(self.path)
                    logger.debug(f"torchaudio.load success, shape={waveform.shape}, sr={sr}")
                    return AudioSamples(data=waveform, sample_rate=sr)
                except Exception as e:
                    logger.error(f"torchaudio.load failed: {e}")
                    raise

            def get_samples_played_in_range(self, start, end):
                # logger.info(f"get_samples_played_in_range {start}-{end}")
                info = self.info
                sr = info.sample_rate
                frame_offset = int(start * sr)
                num_frames = int((end - start) * sr)
                
                waveform, sr = torchaudio.load(
                    self.path,
                    frame_offset=frame_offset,
                    num_frames=num_frames
                )
                return AudioSamples(data=waveform, sample_rate=sr)

        # Inject into pyannote.audio.core.io
        pyannote.audio.core.io.AudioDecoder = AudioDecoder
        pyannote.audio.core.io.AudioStreamMetadata = AudioStreamMetadata
        pyannote.audio.core.io.AudioSamples = AudioSamples
        
        # Inject into pyannote.audio.pipelines.utils (to fix warning)
        try:
            import pyannote.audio.pipelines.utils
            pyannote.audio.pipelines.utils.AudioDecoder = AudioDecoder
            logger.info("Injected AudioDecoder into pyannote.audio.pipelines.utils")
        except Exception as e:
            logger.warning(f"Could not inject into pipelines.utils: {e}")

        logger.info("AudioDecoder monkeypatch applied successfully.")

except Exception as e:
    logger.error(f"Failed to apply AudioDecoder monkeypatch: {e}")

# тяжёлые вещи нужны сразу
import torch
import torchaudio
from scipy.spatial.distance import cdist
from pyannote.core import Annotation

# Pydub – обязательная зависимость
try:
    from pydub import AudioSegment
except ImportError:
    messagebox.showerror(
        "Ошибка",
        "Не установлена библиотека pydub.\n"
        "Установи её командой:\n\n    pip install pydub\n"
    )
    sys.exit(1)

# Google Drive – опционально
try:
    from google.oauth2.credentials import Credentials
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaFileUpload
    from google_auth_oauthlib.flow import InstalledAppFlow
    from google.auth.transport.requests import Request
    GDRIVE_AVAILABLE = True
except ImportError:
    GDRIVE_AVAILABLE = False

# DOCX – тоже опционально
try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# --- КОНСТАНТЫ И НАСТРОЙКИ ---

APP_TITLE = "Нейро Стенографист v0.5 Beta"

DB_FILE = "voice_db.pkl"
SETTINGS_FILE = "settings.json"
RECORDS_DIR = os.path.join(external_path, "Meeting_Records")
SAMPLES_DIR = os.path.join(external_path, "Voice_Samples")
TEMP_DIR = os.path.join(external_path, "temp_segments")

os.makedirs(RECORDS_DIR, exist_ok=True)
os.makedirs(SAMPLES_DIR, exist_ok=True)
GDRIVE_FOLDER = "NeuroStenographer_Records"

BUILTIN_PROMPTS = {
    "Стандартный (умный протокол)": """Твоя роль: профессиональный бизнес-ассистент и редактор.
1. ЗАДАЧА: Исправить ошибки транскрибации по контексту и оформить деловой протокол встречи.
2. ФОРМАТ (Markdown):

# 📝 Протокол встречи от {date}

## 🎯 Краткое содержание
(2–4 абзаца, суть обсуждения и итоги)

## ✅ Принятые решения
* ...

## 🛠 Задачи (Action Items)
| Задача | Ответственный | Срок |
|--------|---------------|------|
| ...    | ...           | ...  |

## ❓ Открытые вопросы
* ...

Тон: деловой, спокойный, без воды.""",

    "IT / Разработка": """Твоя роль: технический менеджер (Tech PM).
Нужно оформить краткий технический отчёт по созвону разработчиков.

Формат (Markdown):

# 🏗️ Tech-созвон {date}

## 📌 Статус и контекст
(какие модули/фичи обсуждали, текущий статус)

## 🚀 Принятые решения
* архитектурные решения
* договорённости по срокам и приоритетам

## 🔧 Задачи
| Задача | Исполнитель | Приоритет |
|--------|-------------|-----------|
| ...    | ...         | ...       |

## 🐞 Риски и проблемы
* баги, техдолг, блокеры""",

    "Продажи / CRM": """Твоя роль: ассистент отдела продаж.
Нужно зафиксировать результаты разговора с клиентом.

Формат (Markdown):

# 💼 Встреча с клиентом {date}

## 👤 Клиент и участники
(кто присутствовал, компания, роли)

## 😫 Боли и потребности
* в чём проблема клиента
* чего они хотят достичь

## 💰 Бюджет и сроки
* примерные цифры и сроки, если обсуждались

## 🛡️ Возражения
* какие сомнения и возражения звучали

## ✅ Следующие шаги
| Наша задача | Действие клиента | Срок контакта |
|------------|------------------|---------------|
| ...        | ...              | ...           |""",

    "Брейншторм": """Твоя роль: креативный скрайбер.
Нужно выловить идеи из хаотичной дискуссии и сохранить энергию обсуждения.

Формат (Markdown):

# 🧠 Брейншторм {date}

## 🌡️ Атмосфера
(как прошла дискуссия, о чём спорили)

## 💎 Основные идеи
* краткие формулировки ключевых идей

## 🗺️ Ход мыслей
(как развивались идеи, какие ветки отбросили)

## 💬 Цитаты
* яркие фразы, можно дословно

## 🚀 К проверке
| Идея | Кто отвечает |
|------|--------------|
| ...  | ...          |"""
}

OLLAMA_MODELS = {
    "Weak (CPU / <8GB VRAM)": [
        {"name": "Llama 3.2 3B", "id": "llama3.2:3b"},
        {"name": "Qwen 2.5 7B", "id": "qwen2.5:7b"},
        {"name": "Gemma 2 9B", "id": "gemma2:9b"},
        {"name": "DeepSeek-R1 7B", "id": "deepseek-r1:7b"},
        {"name": "Mistral 7B", "id": "mistral"},
    ],
    "Medium (16-20GB VRAM)": [
        {"name": "Qwen 2.5 14B", "id": "qwen2.5:14b"},
        {"name": "Mistral Small 24B", "id": "mistral-small"},
        {"name": "Gemma 2 27B", "id": "gemma2:27b"},
        {"name": "DeepSeek-R1 14B", "id": "deepseek-r1:14b"},
        {"name": "GPT-OSS", "id": "gpt-oss"},
    ],
    "Pro (>20GB VRAM)": [
        {"name": "Qwen 2.5 32B", "id": "qwen2.5:32b"},
        {"name": "Llama 3.3 70B", "id": "llama3.3:70b"},
        {"name": "DeepSeek-R1 32B", "id": "deepseek-r1:32b"},
        {"name": "DeepSeek-R1 671B (Distill)", "id": "deepseek-r1:671b"},
    ],
}

DEFAULT_SETTINGS = {
    "hf_token": "",
    "deepgram_key": "",
    "keywords": "",
    "processing_mode": "cloud",          # cloud | local
    "cloud_use_mp3": True,
    "local_model_size": "base",

    # "local_compute": "int8",  # Removed, auto-detected
    "llm_provider": "openrouter",       # openrouter | local
    "or_key": "",
    "or_model": "gpt-4.1-mini",
    "local_model": "qwen2.5:7b",
    "local_url": "http://localhost:11434/v1/chat/completions",
    "current_prompt_name": "Стандартный (умный протокол)",
    "system_prompt": BUILTIN_PROMPTS["Стандартный (умный протокол)"],
    "custom_prompts": {},
    "use_gdrive": False,
    "keep_local": True,
    "input_device": "Default",
    "rec_format": "wav",                # wav | mp3 (итоговый формат)
    "save_txt": True,
    "save_docx": True,
}

ENROLL_TEXT = (
    "Привет! Меня зовут {name}. Я даю согласие на запись моего голоса для создания цифрового слепка. "
    "Сто, двести, триста, четыреста, пятьсот. "
    "В чащах юга жил бы цитрус? Да, но фальшивый экземпляр! "
    "Аэрофотосъёмка ландшафта уже проведена. "
    "Проверка микрофона: раз, два, три. Запись можно завершать."
)

RETRAIN_TEXT = (
    "Это дополнительная запись для улучшения распознавания моего голоса. "
    "Шестьсот, семьсот, восемьсот, девятьсот, тысяча. "
    "Широкая электрификация южных губерний даст мощный толчок подъёму сельского хозяйства. "
    "Эх, чужак, общий съём цен шляп (юфть) — вдрызг! "
    "Теперь система должна узнавать меня намного лучше."
)

ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("blue")


# --- ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ---

def sanitize_filename(name: str) -> str:
    return re.sub(r'[<>:"/\\|?*]', "_", str(name)).strip() or "untitled"


def cosine_distance(e1, e2) -> float:
    if len(e1.shape) == 1:
        e1 = e1.reshape(1, -1)
    if len(e2.shape) == 1:
        e2 = e2.reshape(1, -1)
    return float(cdist(e1, e2, metric="cosine")[0, 0])


def fetch_openrouter_models():
    try:
        resp = requests.get("https://openrouter.ai/api/v1/models")
        if resp.status_code == 200:
            data = resp.json()
            return [m["id"] for m in data.get("data", [])]
    except Exception:
        pass
    return []


# --- МЕНЕДЖЕР НАСТРОЕК ---

class ConfigManager:
    def __init__(self):
        self.data = self._load()

    def _load(self):
        if not os.path.exists(SETTINGS_FILE):
            self._save_data(DEFAULT_SETTINGS)
            return DEFAULT_SETTINGS.copy()
        try:
            with open(SETTINGS_FILE, "r", encoding="utf-8") as f:
                d = json.load(f)
            for k, v in DEFAULT_SETTINGS.items():
                if k not in d:
                    d[k] = v
            return d
        except Exception:
            return DEFAULT_SETTINGS.copy()

    def _save_data(self, d):
        with open(SETTINGS_FILE, "w", encoding="utf-8") as f:
            json.dump(d, f, indent=4, ensure_ascii=False)

    def save(self):
        self._save_data(self.data)

    def get(self, key, default=None):
        return self.data.get(key, default)

    def set(self, key, value):
        self.data[key] = value


config = ConfigManager()


# --- FFmpeg ---

class FFmpegInstaller:
    @staticmethod
    def is_installed():
        try:
            subprocess.run(
                ["ffmpeg", "-version"],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                check=False,
            )
            return True
        except Exception:
            return False

    @staticmethod
    def install(log_cb=print):
        url = "https://www.gyan.dev/ffmpeg/builds/ffmpeg-release-essentials.zip"
        log_cb("Скачивание FFmpeg...")
        try:
            r = requests.get(url, stream=True)
            r.raise_for_status()
        except Exception as e:
            log_cb(f"Ошибка загрузки FFmpeg: {e}")
            return

        with BytesIO(r.content) as b, zipfile.ZipFile(b) as z:
            folder = [n for n in z.namelist() if n.endswith("bin/")][0]
            for f in z.namelist():
                if f.startswith(folder) and f.endswith(".exe"):
                    z.extract(f, ".")
                    src = os.path.join(".", f)
                    dst = os.path.join(".", os.path.basename(f))
                    shutil.move(src, dst)
        log_cb("FFmpeg установлен.")


class AudioHelper:
    @staticmethod
    def get_devices():
        try:
            return [
                f"{i}: {d['name']}"
                for i, d in enumerate(sd.query_devices())
                if d.get("max_input_channels", 0) > 0
            ]
        except Exception:
            return []


# --- РЕКОРДЕР АУДИО ---

class AudioRecorder:
    """
    Пишем всегда во временный WAV-файл, потом по выбору пользователя
    сохраняем как WAV или конвертируем в MP3.
    """

    def __init__(self):
        self.recording = False
        self.queue = queue.Queue()
        self.stream = None
        self.writer_thread = None
        self.temp_filename = None  # всегда WAV
        self.final_target = None   # финальный путь (с расширением пользователя)

    def _callback(self, indata, frames, time_info, status):
        if self.recording:
            self.queue.put(indata.copy())

    def _writer(self):
        with sf.SoundFile(
            self.temp_filename,
            mode="w",
            samplerate=16000,
            channels=1
        ) as file:
            while self.recording or not self.queue.empty():
                try:
                    data = self.queue.get(timeout=0.5)
                    file.write(data)
                except queue.Empty:
                    continue

    def start(self, target_filename: str):
        if self.recording:
            return

        self.final_target = target_filename
        self.temp_filename = f"temp_rec_{uuid.uuid4().hex}.wav"

        dev_conf = config.get("input_device")
        dev_id = None
        if dev_conf and ":" in dev_conf:
            try:
                dev_id = int(dev_conf.split(":")[0])
            except ValueError:
                dev_id = None

        self.recording = True
        try:
            self.stream = sd.InputStream(
                samplerate=16000,
                channels=1,
                callback=self._callback,
                device=dev_id,
            )
            self.stream.start()
        except Exception as e:
            print(f"Ошибка запуска записи: {e}")
            self.recording = False
            return

        self.writer_thread = threading.Thread(target=self._writer, daemon=True)
        self.writer_thread.start()

    def stop(self, force_wav: bool = False) -> bool:
        if not self.recording:
            return False

        self.recording = False

        try:
            if self.stream:
                self.stream.stop()
                self.stream.close()
        except Exception:
            pass

        if self.writer_thread:
            self.writer_thread.join()

        if not self.temp_filename or not os.path.exists(self.temp_filename):
            return False

        final_fmt = config.get("rec_format")
        if force_wav:
            final_fmt = "wav"

        try:
            dirpath = os.path.dirname(self.final_target)
            if dirpath:
                os.makedirs(dirpath, exist_ok=True)

            if final_fmt == "mp3":
                # если пользователь указал .wav, заменим на .mp3
                if self.final_target.endswith(".wav"):
                    self.final_target = self.final_target[:-4] + ".mp3"

                AudioSegment.from_wav(self.temp_filename).export(
                    self.final_target,
                    format="mp3",
                    bitrate="128k",
                )
                os.remove(self.temp_filename)
            else:
                if os.path.exists(self.final_target):
                    os.remove(self.final_target)
                shutil.move(self.temp_filename, self.final_target)

            return True
        except Exception as e:
            print(f"Ошибка сохранения аудио: {e}")
            return False


# --- AI БЛОК: STT + ДИАРИЗАЦИЯ ---

class AIProcessor:
    def __init__(self):
        self.device = None
        self._emb_model = None
        self._emb_inference = None
        self._pipeline = None
        self._whisper_model = None

    def _init_torch(self):
        if self.device:
            return
        self.device = "cuda" if torch.cuda.is_available() else "cpu"

    def _load_embedding_model(self):
        if self._emb_inference:
            return
        self._init_torch()
        from pyannote.audio import Model, Inference

        token = config.get("hf_token")
        self._emb_model = Model.from_pretrained(
            "pyannote/wespeaker-voxceleb-resnet34-LM",
            token=token,
        )
        if self.device == "cuda":
            self._emb_model.to(torch.device("cuda"))
        self._emb_inference = Inference(self._emb_model, window="whole")

    def _load_pipeline(self):
        if self._pipeline:
            return
        self._init_torch()
        # AudioDecoder is already injected by global monkeypatch
        
        from pyannote.audio import Pipeline

        token = config.get("hf_token")
        self._pipeline = Pipeline.from_pretrained(
            "pyannote/speaker-diarization-3.1",
            token=token,
        )
        self._pipeline.to(torch.device(self.device))

    def _load_whisper(self):
        if self._whisper_model:
            return
        self._init_torch()
        from faster_whisper import WhisperModel

        self._whisper_model = WhisperModel(
            config.get("local_model_size"),
            device=self.device,
            compute_type="float16" if self.device == "cuda" else "int8",
        )

    def create_embedding(self, wav_path: str):
        self._load_embedding_model()
        wave, sr = torchaudio.load(wav_path)
        return self._emb_inference({"waveform": wave, "sample_rate": sr})

    def analyze(self, path: str, voice_db: dict, log_cb):
        mode = config.get("processing_mode")
        if mode == "cloud":
            return self._analyze_cloud(path, voice_db, log_cb)
        return self._analyze_local(path, voice_db, log_cb)

    # --- Облако (Deepgram) ---

    def _analyze_cloud(self, path: str, voice_db: dict, log_cb):
        log_cb("Обработка в облаке (Deepgram)...")

        key = config.get("deepgram_key")
        if not key:
            raise RuntimeError("Не задан ключ Deepgram API.")

        src_path = path
        tmp_mp3 = f"tmp_{uuid.uuid4().hex}.mp3"

        if config.get("cloud_use_mp3"):
            log_cb("Сжатие аудио в MP3 для отправки...")
            AudioSegment.from_file(path).export(
                tmp_mp3,
                format="mp3",
                bitrate="64k",
            )
            src_path = tmp_mp3

        url = (
            "https://api.deepgram.com/v1/listen"
            "?model=nova-2"
            "&diarize=true"
            "&smart_format=true"
            "&language=ru"
            "&punctuate=true"
        )

        keywords = config.get("keywords")
        if keywords:
            for w in keywords.split(","):
                w = w.strip()
                if w:
                    url += f"&keywords={w}:2"

        try:
            with open(src_path, "rb") as f:
                r = requests.post(
                    url,
                    headers={"Authorization": f"Token {key}"},
                    data=f,
                )
        except Exception as e:
            raise ConnectionError(f"Сетевая ошибка Deepgram: {e}")

        if src_path == tmp_mp3 and os.path.exists(tmp_mp3):
            os.remove(tmp_mp3)

        if r.status_code != 200:
            raise RuntimeError(f"Ошибка Deepgram: {r.text}")

        words = (
            r.json()
            .get("results", {})
            .get("channels", [{}])[0]
            .get("alternatives", [{}])[0]
            .get("words", [])
        )

        segments = []
        current = {"speaker": None, "words": [], "start": 0.0, "end": 0.0}

        for w in words:
            spk = w.get("speaker", 0)
            word = w.get("punctuated_word") or w.get("word", "")
            if (
                current["speaker"] is None
                or spk == current["speaker"]
            ) and current["words"]:
                # продолжаем текущий сегмент
                current["words"].append(word)
                current["end"] = w["end"]
            else:
                # новый спикер или первая фраза
                if current["words"]:
                    segments.append(
                        {
                            "start": current["start"],
                            "end": current["end"],
                            "label": f"Спикер {current['speaker']}",
                            "text": " ".join(current["words"]),
                            "audio": None,
                            "fs": 0,
                        }
                    )
                current = {
                    "speaker": spk,
                    "words": [word],
                    "start": w["start"],
                    "end": w["end"],
                }

        if current["words"]:
            segments.append(
                {
                    "start": current["start"],
                    "end": current["end"],
                    "label": f"Спикер {current['speaker']}",
                    "text": " ".join(current["words"]),
                    "audio": None,
                    "fs": 0,
                }
            )

        return segments, {}  # вооблаке идентификацию по базе не делаем

    # --- Локальный режим (PyAnnote + faster-whisper) ---

    def _analyze_local(self, path: str, voice_db: dict, log_cb):
        log_cb("Локальный режим: диаризация...")
        self._load_pipeline()
        diar = self._pipeline(path)
        # logger.debug(f"Pipeline returned type: {type(diar)}")

        if isinstance(diar, Annotation):
            ann = diar
        elif hasattr(diar, "annotation"):
            ann = diar.annotation
        elif hasattr(diar, "speaker_diarization"):
            # Handle DiarizeOutput from newer pyannote versions
            ann = diar.speaker_diarization
        else:
            raise RuntimeError(f"Пайплайн вернул неизвестный тип: {type(diar)}")

        audio = AudioSegment.from_file(path)
        speaker_chunks = []

        os.makedirs(TEMP_DIR, exist_ok=True)

        for seg, _, label in ann.itertracks(yield_label=True):
            start_ms = int(seg.start * 1000)
            end_ms = int(seg.end * 1000)
            if end_ms - start_ms < 500:
                continue
            chunk = audio[start_ms:end_ms]
            tmp_wav = os.path.join(TEMP_DIR, f"chunk_{uuid.uuid4().hex}.wav")
            chunk.export(tmp_wav, format="wav")
            speaker_chunks.append(
                {
                    "start": seg.start,
                    "end": seg.end,
                    "label": label,
                    "file": tmp_wav,
                }
            )

        if not speaker_chunks:
            return [], {}

        log_cb(f"Локальный режим: распознавание речи ({len(speaker_chunks)} фрагментов)...")
        self._load_whisper()

        segments = []
        batch_size = config.get("batch_size", 8)
        total_chunks = len(speaker_chunks)
        
        # Process chunks in batches
        for batch_idx in range(0, total_chunks, batch_size):
            batch = speaker_chunks[batch_idx:batch_idx + batch_size]
            batch_num = (batch_idx // batch_size) + 1
            total_batches = (total_chunks + batch_size - 1) // batch_size
            
            log_cb(f"Обработка пакета {batch_num}/{total_batches} ({len(batch)} фрагментов)...")
            
            for ch in batch:
                wav, sr = torchaudio.load(ch["file"])
                if sr != 16000:
                    wav = torchaudio.functional.resample(wav, sr, 16000)
                    sr = 16000
                # Rewrite the file to ensure correct format for Whisper
                torchaudio.save(ch["file"], wav, sr)

                try:
                    res, info = self._whisper_model.transcribe(
                        ch["file"], beam_size=5
                    )
                except RuntimeError as e:
                    if "cublas" in str(e).lower() and self.device == "cuda":
                        logger.warning("CUDA error detected (missing libraries?). Falling back to CPU.")
                        log_cb("Ошибка CUDA. Переключение на CPU...")
                        self.device = "cpu"
                        self._whisper_model = None
                        self._load_whisper()
                        # Retry with CPU model
                        res, info = self._whisper_model.transcribe(
                            ch["file"], beam_size=5
                        )
                    else:
                        raise e
                text = " ".join(r.text.strip() for r in res)
                segments.append(
                    {
                        "start": ch["start"],
                        "end": ch["end"],
                        "label": ch["label"],
                        "text": text,
                        "audio": ch["file"],
                        "fs": 16000,
                    }
                )

        # подготовка неизвестных спикеров для мастера
        unknown = {}
        for s in segments:
            label = s["label"]
            if label in voice_db:
                continue
            if label not in unknown:
                unknown[label] = {
                    "audio": s["audio"],
                    "fs": s["fs"],
                    "dur": s["end"] - s["start"],
                }
            else:
                if s["end"] - s["start"] > unknown[label]["dur"]:
                    unknown[label]["audio"] = s["audio"]
                    unknown[label]["fs"] = s["fs"]
                    unknown[label]["dur"] = s["end"] - s["start"]

        return segments, unknown

    def check_models_integrity(self, log_cb):
        log_cb("Проверка целостности моделей...")
        logger.info("Starting integrity check")
        
        # 1. Check CUDA
        cuda_ok = torch.cuda.is_available()
        log_cb(f"CUDA доступна: {'Да' if cuda_ok else 'Нет'}")
        logger.info(f"CUDA available: {cuda_ok}")
        
        if not cuda_ok:
            log_cb("")
            log_cb("⚠️ CUDA недоступна. Возможные причины:")
            log_cb("1. Не установлен PyTorch с поддержкой CUDA")
            log_cb("2. Установите: pip install torch torchvision torchaudio --index-url https://download.pytorch.org/whl/cu118")
            log_cb("3. Или используйте облачный режим (Deepgram)")
            logger.warning("CUDA not available - likely PyTorch CPU-only version")
        
        # 2. Check Whisper
        try:
            log_cb("Загрузка Whisper (тест)...")
            self._load_whisper()
            log_cb("Whisper загружен успешно.")
            logger.info("Whisper loaded successfully")
        except Exception as e:
            log_cb(f"Ошибка Whisper: {e}")
            logger.error(f"Whisper error: {e}", exc_info=True)

        # 3. Check Pyannote
        try:
            log_cb("Загрузка Pyannote (тест)...")
            self._load_pipeline()
            log_cb("Pyannote загружен успешно.")
            logger.info("Pyannote loaded successfully")
        except Exception as e:
            log_cb(f"Ошибка Pyannote: {e}")
            logger.error(f"Pyannote error: {e}", exc_info=True)
            
        log_cb("Проверка завершена.")
        logger.info("Integrity check completed")


# --- GOOGLE DRIVE ---

class GDriveClient:
    def __init__(self, log_cb=print):
        self.log_cb = log_cb
        self.ok = GDRIVE_AVAILABLE
        self.creds = None
        self.svc = None
        self.folder_id = None

    def log(self, msg: str):
        self.log_cb(f"[Google Drive] {msg}")

    def auth(self) -> bool:
        if not self.ok:
            return False

        scopes = ["https://www.googleapis.com/auth/drive.file"]

        if os.path.exists("token.json"):
            try:
                self.creds = Credentials.from_authorized_user_file(
                    "token.json", scopes
                )
            except Exception:
                os.remove("token.json")
                self.creds = None

        if not self.creds or not self.creds.valid:
            if self.creds and self.creds.expired:
                try:
                    self.creds.refresh(Request())
                except Exception:
                    self.creds = None

            if not self.creds and os.path.exists("credentials.json"):
                try:
                    flow = InstalledAppFlow.from_client_secrets_file(
                        "credentials.json", scopes
                    )
                    self.creds = flow.run_local_server(port=0)
                except Exception:
                    return False

            if self.creds:
                with open("token.json", "w") as token:
                    token.write(self.creds.to_json())

        if not self.creds:
            return False

        try:
            self.svc = build("drive", "v3", credentials=self.creds)
            resp = (
                self.svc.files()
                .list(
                    q=f"name='{GDRIVE_FOLDER}' and mimeType='application/vnd.google-apps.folder' and trashed=false",
                    spaces="drive",
                )
                .execute()
            )
            files = resp.get("files", [])
            if files:
                self.folder_id = files[0]["id"]
            else:
                meta = {
                    "name": GDRIVE_FOLDER,
                    "mimeType": "application/vnd.google-apps.folder",
                }
                folder = (
                    self.svc.files()
                    .create(body=meta, fields="id")
                    .execute()
                )
                self.folder_id = folder["id"]
            return True
        except Exception:
            return False

    def upload(self, path: str):
        if not os.path.exists(path):
            self.log(f"Файл не найден, пропускаю загрузку: {path}")
            return

        if not self.auth():
            self.log("Не удалось авторизоваться в Google Drive.")
            return

        try:
            file_metadata = {
                "name": os.path.basename(path),
                "parents": [self.folder_id],
            }
            media = MediaFileUpload(path, resumable=True)
            self.svc.files().create(
                body=file_metadata, media_body=media
            ).execute()
            self.log(f"Загружен файл: {os.path.basename(path)}")
        except Exception as e:
            self.log(f"Ошибка загрузки: {e}")


# --- ГЕНЕРАЦИЯ DOCX ИЗ MARKDOWN-ПОДОБНОГО ТЕКСТА ---

class DocxGenerator:
    @staticmethod
    def create_report(markdown_text: str, filename: str):
        if not DOCX_AVAILABLE:
            # резервный вариант – просто сохранить .md/.txt
            alt = filename.replace(".docx", ".md")
            with open(alt, "w", encoding="utf-8") as f:
                f.write(markdown_text)
            return

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        lines = markdown_text.splitlines()
        table_data = []
        in_table = False

        for line in lines:
            if line.startswith("|") and line.endswith("|"):
                cols = [c.strip() for c in line.strip("|").split("|")]
                table_data.append(cols)
                in_table = True
                continue
            else:
                if in_table:
                    DocxGenerator._render_table(doc, table_data)
                    table_data = []
                    in_table = False

            if line.startswith("# "):
                doc.add_heading(line[2:], level=0)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=1)
            elif line.startswith("* "):
                p = doc.add_paragraph(style="List Bullet")
                DocxGenerator._render_runs(p, line[2:])
            else:
                p = doc.add_paragraph()
                DocxGenerator._render_runs(p, line)

        if in_table:
            DocxGenerator._render_table(doc, table_data)

        doc.save(filename)

    @staticmethod
    def _render_table(doc, table_data):
        if not table_data:
            return
        rows = len(table_data)
        cols = len(table_data[0])
        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"

        for r, row in enumerate(table_data):
            for c, val in enumerate(row):
                cell = table.rows[r].cells[c]
                cell.text = val
                if r == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

    @staticmethod
    def _render_runs(paragraph, text: str):
        # очень простой парсер **жирного** текста
        parts = re.split(r"(\*\*.*?\*\*)", text)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)


# --- LLM-КЛИЕНТ ---

class LLMClient:
    def summarize(self, transcript_text: str) -> str:
        prompt_template = config.get("system_prompt") or ""
        prompt = prompt_template.replace(
            "{date}", datetime.now().strftime("%d.%m.%Y")
        )

        messages = [
            {"role": "system", "content": prompt},
            {"role": "user", "content": transcript_text},
        ]

        provider = config.get("llm_provider")

        if provider == "openrouter":
            url = "https://openrouter.ai/api/v1/chat/completions"
            headers = {
                "Authorization": f"Bearer {config.get('or_key')}",
                "Content-Type": "application/json",
            }
            data = {
                "model": config.get("or_model"),
                "messages": messages,
            }
        else:
            # локальный OpenAI-совместимый сервер (например, прокси к Ollama)
            url = config.get("local_url")
            headers = {"Content-Type": "application/json"}
            data = {
                "model": config.get("local_model"),
                "messages": messages,
                "stream": False,
            }

        try:
            resp = requests.post(url, headers=headers, json=data)
            if resp.status_code != 200:
                return f"Ошибка LLM: {resp.status_code} {resp.text}"
            j = resp.json()
            return j["choices"][0]["message"]["content"]
        except Exception as e:
            return f"Ошибка LLM: {e}"


# --- OLLAMA ---

class OllamaManager:
    @staticmethod
    def is_installed() -> bool:
        try:
            r = subprocess.run(
                ["ollama", "list"],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
            )
            return r.returncode == 0
        except Exception:
            return False

    @staticmethod
    def get_local_models():
        if not OllamaManager.is_installed():
            return []
        r = subprocess.run(
            ["ollama", "list"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
        lines = [l.strip() for l in r.stdout.splitlines() if l.strip()]
        models = []
        for line in lines[1:]:
            parts = line.split()
            if parts:
                models.append(parts[0])
        return models

    @staticmethod
    def pull_model(model_id: str, progress_cb, done_cb):
        """
        Запускаем `ollama pull model_id` и парсим проценты из stdout.
        progress_cb(progress: float, message: str)
        done_cb(success: bool, model_id: str)
        """
        try:
            proc = subprocess.Popen(
                ["ollama", "pull", model_id],
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                text=True,
            )
        except Exception as e:
            done_cb(False, model_id)
            return

        for line in proc.stdout:
            m = re.search(r"(\d+)%", line)
            if m:
                pct = int(m.group(1))
                progress_cb(pct / 100.0, f"Загрузка {model_id}: {pct}%")

        proc.wait()
        done_cb(proc.returncode == 0, model_id)


class OllamaModelShop(ctk.CTkToplevel):
    def __init__(self, parent, model_var: ctk.StringVar):
        super().__init__(parent)
        self.title("Магазин моделей Ollama")
        self.geometry("520x480")
        self.attributes("-topmost", True)  # Always on top
        self.model_var = model_var
        
        # Get list of installed models once
        self.installed_models = set(OllamaManager.get_local_models() or [])

        self.progress = ctk.CTkProgressBar(self, mode="determinate")
        self.progress.pack(fill="x", padx=10, pady=(10, 0))
        self.progress.set(0.0)

        self.status_label = ctk.CTkLabel(self, text="")
        self.status_label.pack(padx=10, pady=(0, 10))

        self._build_ui()

    def _build_ui(self):
        tabs = ctk.CTkTabview(self)
        tabs.pack(fill="both", expand=True, padx=10, pady=10)

        for group_name, models in OLLAMA_MODELS.items():
            frame = tabs.add(group_name)
            for m in models:
                model_id = m["id"]
                is_installed = model_id in self.installed_models
                
                row = ctk.CTkFrame(frame)
                row.pack(fill="x", padx=5, pady=5)
                
                # Model name label with optional installed indicator
                name_text = f"✓ {m['name']}" if is_installed else m["name"]
                label_color = ("#2ecc71", "#27ae60") if is_installed else ("gray10", "gray90")
                ctk.CTkLabel(
                    row, 
                    text=name_text,
                    text_color=label_color
                ).pack(side="left", padx=5)
                
                # Select button (disabled if not installed)
                select_btn = ctk.CTkButton(
                    row,
                    text="Выбрать",
                    command=lambda mid=model_id: self._select_model(mid),
                    state="normal" if is_installed else "disabled"
                )
                select_btn.pack(side="right", padx=5)
                
                # Download button (changes to Re-download if installed)
                download_text = "Обновить" if is_installed else "Скачать"
                ctk.CTkButton(
                    row,
                    text=download_text,
                    command=lambda mid=model_id: self._pull_model(mid),
                ).pack(side="right", padx=5)

    def log(self, text: str):
        self.status_label.configure(text=text)

    def _select_model(self, model_id: str):
        self.model_var.set(model_id)
        self.status_label.configure(text=f"Выбрана модель: {model_id}")

    def _pull_model(self, model_id: str):
        if not OllamaManager.is_installed():
            messagebox.showerror(
                "Ollama",
                "Ollama не установлена или не найдена в PATH.",
            )
            return

        self.progress.set(0.0)
        self.status_label.configure(text=f"Загрузка модели {model_id}...")
        thread = threading.Thread(
            target=self._pull_thread, args=(model_id,), daemon=True
        )
        thread.start()

    def _pull_thread(self, model_id: str):
        def on_progress(frac, msg):
            self.after(
                0,
                lambda: (self.progress.set(frac), self.status_label.configure(text=msg)),
            )

        def on_done(success, mid):
            def ui():
                if success:
                    self.progress.set(1.0)
                    self.status_label.configure(
                        text=f"Модель {mid} успешно загружена."
                    )
                    # обновляем список доступных моделей
                    models = OllamaManager.get_local_models() or ["None"]
                    if mid in models:
                        self.model_var.set(mid)
                    # Refresh installed models list and rebuild UI
                    self.installed_models = set(models)
                    self._rebuild_ui()
                else:
                    self.status_label.configure(
                        text=f"Ошибка загрузки модели {mid}."
                    )

            self.after(0, ui)

        OllamaManager.pull_model(model_id, on_progress, on_done)
    
    def _rebuild_ui(self):
        """Rebuild the UI after downloading a model to update installed indicators"""
        # Find and destroy the tabs widget
        for widget in self.winfo_children():
            if isinstance(widget, ctk.CTkTabview):
                widget.destroy()
                break
        
        # Rebuild tabs with updated installed models
        self._build_ui()




# --- МАСТЕР ОПРЕДЕЛЕНИЯ НЕИЗВЕСТНЫХ СПИКЕРОВ ---

class IdentifyWizard(ctk.CTkToplevel):
    def __init__(self, parent, unknown_map: dict, voice_db: dict):
        super().__init__(parent)
        self.title("Определение спикеров")
        self.geometry("420x320")
        self.unknown = unknown_map
        self.voice_db = voice_db
        self.result_names = {}
        self.save_flags = {}
        self.keys = list(unknown_map.keys())
        self.index = 0
        
        # Collect all available names: from DB + any new ones assigned in this session
        self.available_names = sorted(list(voice_db.keys()))

        self._build_ui()

    def _build_ui(self):
        self.lift()
        self.focus_force()
        self.grab_set()

        self.label = ctk.CTkLabel(self, text="Неизвестный спикер")
        self.label.pack(pady=15)

        ctk.CTkButton(
            self,
            text="▶️ Прослушать пример",
            command=self._play_sample,
        ).pack(pady=5)

        ctk.CTkLabel(self, text="Имя спикера:").pack(pady=(15, 5))
        
        # Use ComboBox instead of Entry
        self.name_combo = ctk.CTkComboBox(self, values=self.available_names)
        self.name_combo.set("")
        self.name_combo.pack(pady=5)

        self.save_var = ctk.BooleanVar(value=True)
        ctk.CTkCheckBox(
            self, text="Сохранить в базу голосов", variable=self.save_var
        ).pack(pady=5)

        ctk.CTkButton(self, text="Далее", command=self._next).pack(
            pady=15
        )

        self._update_ui()

    def _update_ui(self):
        if self.index >= len(self.keys):
            self.destroy()
            return
        spk_id = self.keys[self.index]
        self.label.configure(text=f"Фрагмент спикера: {spk_id}")
        self.name_combo.set("")
        self.save_var.set(True)
        
        # Update available names in case user added a new one in previous step
        current_session_names = sorted(list(set(self.result_names.values())))
        all_names = sorted(list(set(list(self.voice_db.keys()) + current_session_names)))
        self.name_combo.configure(values=all_names)

    def _play_sample(self):
        if self.index >= len(self.keys):
            return
        spk_id = self.keys[self.index]
        sample = self.unknown[spk_id]
        try:
            data, sr = sf.read(sample["audio"])
            sd.play(data, sample["fs"])
        except Exception:
            pass

    def _next(self):
        if self.index < len(self.keys):
            name = self.name_combo.get().strip()
            spk_id = self.keys[self.index]
            if name:
                self.result_names[spk_id] = name
                self.save_flags[spk_id] = self.save_var.get()
                
                # Auto-save sample to speaker folder
                import shutil
                import glob
                from pydub import AudioSegment
                
                speaker_dir = os.path.join(SAMPLES_DIR, sanitize_filename(name))
                os.makedirs(speaker_dir, exist_ok=True)
                
                # Find next sample number
                existing_samples = glob.glob(os.path.join(speaker_dir, "sample_*.wav"))
                next_num = len(existing_samples) + 1
                sample_path = os.path.join(speaker_dir, f"sample_{next_num:03d}.wav")
                
                # Convert audio to WAV format if needed
                ch = self.unknown[spk_id]
                source_file = ch["audio"]
                
                # Check if already WAV
                if source_file.lower().endswith('.wav'):
                    shutil.copy(source_file, sample_path)
                else:
                    # Convert to WAV using pydub
                    try:
                        audio = AudioSegment.from_file(source_file)
                        audio.export(sample_path, format="wav")
                        logger.info(f"Converted {os.path.basename(source_file)} to WAV")
                    except Exception as e:
                        logger.warning(f"Failed to convert {source_file}: {e}. Copying as-is.")
                        shutil.copy(source_file, sample_path)
                
                logger.info(f"Saved sample for {name}: {sample_path}")
                
                # Update or create speaker in voice_db with initial status
                if name not in self.voice_db:
                    self.voice_db[name] = {
                        "embedding": None,
                        "trained": False,
                        "sample_count": next_num
                    }
                else:
                    # Update sample count for existing speaker
                    if isinstance(self.voice_db[name], dict):
                        self.voice_db[name]["sample_count"] = next_num
        self.index += 1
        self._update_ui()


# --- ДОБАВЛЕНИЕ СПИКЕРА ---

class AddSpeakerDialog(ctk.CTkToplevel):
    def __init__(self, master, voice_db, recorder, ai, log_cb, on_done, existing_name=None):
        super().__init__(master)
        self.title("Добавление голоса" if not existing_name else f"Дообучение: {existing_name}")
        self.geometry("500x650")
        self.voice_db = voice_db
        self.recorder = recorder
        self.ai = ai
        self.log_cb = log_cb
        self.on_done = on_done
        self.existing_name = existing_name

        self._build_ui()

    def _build_ui(self):
        self.lift()
        self.focus_force()
        self.grab_set()

        ctk.CTkLabel(self, text="Имя участника:", font=("Segoe UI", 14, "bold")).pack(pady=(20, 5))
        self.name_entry = ctk.CTkEntry(self, width=300)
        self.name_entry.pack(pady=5)
        
        # If re-training existing speaker, pre-fill and disable name field
        if self.existing_name:
            self.name_entry.insert(0, self.existing_name)
            self.name_entry.configure(state="disabled")

        # Instruction text box (resizable)
        instruction_frame = ctk.CTkFrame(self, fg_color="transparent")
        instruction_frame.pack(pady=10, padx=20, fill="both", expand=True)
        
        instruction_text = (
            "Для обучения голоса произнесите текст ниже:\n\n"
            f"{ENROLL_TEXT.replace('{name}', '[имя]')}"
        ) if not self.existing_name else (
            "Запишите дополнительный семпл голоса:\n\n"
            f"{RETRAIN_TEXT}"
        )
        
        self.instruction_box = ctk.CTkTextbox(
            instruction_frame,
            font=("Segoe UI", 12),
            wrap="word",
            height=150
        )
        self.instruction_box.pack(fill="both", expand=True)
        self.instruction_box.insert("0.0", instruction_text)
        self.instruction_box.configure(state="disabled")

        ctk.CTkLabel(self, text="Запись с микрофона:", font=("Segoe UI", 12, "bold")).pack(pady=(15, 5))
        ctk.CTkButton(
            self,
            text="🎤 Начать запись",
            command=self._start_rec,
            height=35
        ).pack(pady=5)
        self.stop_btn = ctk.CTkButton(
            self,
            text="⏹ Остановить запись",
            fg_color="red",
            state="disabled",
            command=self._stop_rec,
            height=35
        )
        self.stop_btn.pack(pady=5)

        ctk.CTkLabel(self, text="Или загрузить готовый файл:", font=("Segoe UI", 12, "bold")).pack(
            pady=(15, 5)
        )
        ctk.CTkButton(
            self, 
            text="📂 Выбрать файл...", 
            command=self._upload_file,
            height=35
        ).pack(pady=5)

        self.timer_label = ctk.CTkLabel(self, text="00:00", font=("Segoe UI", 20, "bold"))
        self.timer_label.pack(pady=5)

        self.progress = ctk.CTkProgressBar(self, mode="indeterminate")
        # pack only when processing
        
        self.status_label = ctk.CTkLabel(self, text="", text_color="gray")
        self.status_label.pack(pady=5)

        self.rec_start_time = 0
        self.is_recording = False

    def _start_rec(self):
        filename = f"enroll_{uuid.uuid4().hex}.wav"
        self.recorder.start(filename)
        self.stop_btn.configure(state="normal")
        self.is_recording = True
        self.rec_start_time = time.time()
        self._update_timer()

    def _update_timer(self):
        if self.is_recording:
            elapsed = int(time.time() - self.rec_start_time)
            self.timer_label.configure(text=f"{elapsed // 60:02}:{elapsed % 60:02}")
            self.after(100, self._update_timer)

    def _stop_rec(self):
        self.is_recording = False
        ok = self.recorder.stop(force_wav=True)
        self.stop_btn.configure(state="disabled")
        if ok:
            self._start_processing(self.recorder.final_target)

    def _upload_file(self):
        path = filedialog.askopenfilename()
        if path:
            self._start_processing(path)

    def _start_processing(self, path):
        name = self.name_entry.get().strip()
        if not name:
            messagebox.showwarning("Имя", "Сначала введи имя участника.")
            return

        self.progress.pack(pady=10)
        self.progress.start()
        self.status_label.configure(text="Идёт обработка... (Может занять 10-20 сек)")
        
        threading.Thread(
            target=self._process_sample_thread,
            args=(path, name),
            daemon=True
        ).start()

    def _process_sample_thread(self, path, name):
        try:
            import glob
            from pydub import AudioSegment
            
            # Create speaker-specific directory
            speaker_dir = os.path.join(SAMPLES_DIR, sanitize_filename(name))
            os.makedirs(speaker_dir, exist_ok=True)
            
            # Find next sample number
            existing_samples = glob.glob(os.path.join(speaker_dir, "sample_*.wav"))
            next_num = len(existing_samples) + 1
            sample_path = os.path.join(speaker_dir, f"sample_{next_num:03d}.wav")
            
            self.log_cb(f"Обучение голоса: {name}")
            
            # Convert to WAV and save
            AudioSegment.from_file(path).export(sample_path, format="wav")
            
            # Create embedding
            emb = self.ai.create_embedding(sample_path)
            
            # Update voice_db with new format
            self.voice_db[name] = {
                "embedding": emb,
                "trained": True,
                "sample_count": next_num
            }
            
            self.after(0, self.on_done)
            self.after(0, lambda: messagebox.showinfo("Успех", "Голос успешно добавлен!"))
            self.after(0, self.destroy)
        except Exception as e:
            self.log_cb(f"Ошибка обучения: {e}")
            self.after(0, lambda: self.progress.stop())
            self.after(0, lambda: self.progress.pack_forget())
            self.after(0, lambda: self.status_label.configure(text=f"Ошибка: {e}"))


# --- ОСНОВНОЕ ОКНО ПРИЛОЖЕНИЯ ---

import pickle


class App(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title(APP_TITLE)
        self.geometry("1200x800")



        self.voice_db = (
            pickle.load(open(DB_FILE, "rb"))
            if os.path.exists(DB_FILE)
            else {}
        )
    
        # Migrate old voice_db format to new format
        self._migrate_voice_db()

        self.recorder = AudioRecorder()
        self.ai = AIProcessor()
        self.llm = LLMClient()
        self.gdrive = GDriveClient(self._log)

        self.is_recording = False
        self.record_start_time = 0
        self.last_transcript_text = ""
        self.last_basename = ""

        self._build_ui()

        if not config.get("hf_token"):
            self.after(
                1500,
                lambda: messagebox.showwarning(
                    "HF Token",
                    "Не задан HF Token. Идентификация голосов работать не будет.",
                ),
            )

    def _get_host_dirs(self, host_name: str):
        """
        Returns paths for (recordings_dir, transcripts_dir, reports_dir)
        based on the host name. Creates them if they don't exist.
        Structure:
          Meeting_Records/
            [Host Name]/
              Записи/
              Транскрипции/
              Отчеты/
        """
        safe_host = sanitize_filename(host_name)
        base_dir = os.path.join(RECORDS_DIR, safe_host)
        
        rec_dir = os.path.join(base_dir, "Записи")
        trans_dir = os.path.join(base_dir, "Транскрипции")
        rep_dir = os.path.join(base_dir, "Отчеты")
        
        os.makedirs(rec_dir, exist_ok=True)
        os.makedirs(trans_dir, exist_ok=True)
        os.makedirs(rep_dir, exist_ok=True)
        
        return rec_dir, trans_dir, rep_dir

    def _migrate_voice_db(self):
        """Migrate old voice_db format (name->embedding) to new format (name->dict)"""
        import glob
        
        migrated = False
        for name, data in list(self.voice_db.items()):
            if isinstance(data, np.ndarray):
                # Old format - convert to new format
                speaker_dir = os.path.join(SAMPLES_DIR, sanitize_filename(name))
                sample_count = len(glob.glob(os.path.join(speaker_dir, "sample_*.wav"))) if os.path.exists(speaker_dir) else 0
                
                self.voice_db[name] = {
                    "embedding": data,
                    "trained": True,
                    "sample_count": sample_count
                }
                migrated = True
                logger.info(f"Migrated {name} to new format (sample_count: {sample_count})")
            elif isinstance(data, dict):
                # New format - ensure all fields exist
                if "trained" not in data:
                    data["trained"] = data.get("embedding") is not None
                if "sample_count" not in data:
                    speaker_dir = os.path.join(SAMPLES_DIR, sanitize_filename(name))
                    data["sample_count"] = len(glob.glob(os.path.join(speaker_dir, "sample_*.wav"))) if os.path.exists(speaker_dir) else 0
        
        if migrated:
            self._save_db()
            logger.info("Voice DB migration completed")
    
    def _get_speaker_sample_count(self, name):
        """Get the number of samples for a speaker"""
        import glob
        speaker_dir = os.path.join(SAMPLES_DIR, sanitize_filename(name))
        if not os.path.exists(speaker_dir):
            return 0
        return len(glob.glob(os.path.join(speaker_dir, "sample_*.wav")))

    # --- служебные методы ---

    def _save_db(self):
        pickle.dump(self.voice_db, open(DB_FILE, "wb"))

    def _log(self, msg: str):
        """Log message to status bar and log_box if it exists"""
        ts = datetime.now().strftime("%H:%M:%S")
        formatted_msg = f"[{ts}] {msg}"
        
        # Update status bar
        if hasattr(self, 'status_label'):
            self.status_label.configure(text=msg)
        
        # Also log to textbox if it exists (analysis page)
        if hasattr(self, 'log_box') and self.log_box.winfo_exists():
            try:
                self.log_box.insert("end", formatted_msg + "\n")
                self.log_box.see("end")
            except:
                pass

    # --- UI ---

    def _build_ui(self):
        """Build the main UI with sidebar navigation"""
        # Main layout: sidebar (left) + content (right)
        self.grid_columnconfigure(0, weight=0)  # Sidebar - fixed width
        self.grid_columnconfigure(1, weight=1)  # Content - expandable
        self.grid_rowconfigure(0, weight=1)
        
        # === SIDEBAR ===
        self.sidebar = ctk.CTkFrame(self, width=180, corner_radius=0, fg_color=("#dbdbdb", "#2b2b2b"))
        self.sidebar.grid(row=0, column=0, sticky="nsew")
        self.sidebar.grid_propagate(False)
        
        # App title in sidebar
        ctk.CTkLabel(
            self.sidebar, 
            text="Meeting App", 
            font=("Segoe UI", 16, "bold")
        ).pack(pady=20)
        
        # Navigation buttons
        self.nav_buttons = {}
        nav_items = [
            ("🎤 Запись", self._show_recording_page),
            ("📊 Анализ", self._show_analysis_page),
            ("🗣️ Голоса", self._show_voices_page),
            ("⚙️ Настройки", self._show_settings_page),
            ("📁 Файлы", self._show_files_page),
        ]
        
        for text, command in nav_items:
            btn = ctk.CTkButton(
                self.sidebar,
                text=text,
                command=command,
                width=160,
                height=40,
                anchor="w",
                fg_color="transparent",
                text_color=("gray10", "gray90"),
                hover_color=("gray70", "gray30")
            )
            btn.pack(pady=5, padx=10)
            self.nav_buttons[text] = btn
        
        # === CONTENT AREA ===
        self.content_frame = ctk.CTkFrame(self, corner_radius=0)
        self.content_frame.grid(row=0, column=1, sticky="nsew", padx=0, pady=0)
        self.content_frame.grid_columnconfigure(0, weight=1)
        self.content_frame.grid_rowconfigure(0, weight=1)
        
        # === STATUS BAR (bottom) ===
        self.status_bar = ctk.CTkFrame(self, height=30, corner_radius=0)
        self.status_bar.grid(row=1, column=0, columnspan=2, sticky="ew")
        
        self.status_label = ctk.CTkLabel(
            self.status_bar, 
            text="Готов", 
            anchor="w"
        )
        self.status_label.pack(side="left", padx=10)
        
        self.progress = ctk.CTkProgressBar(self.status_bar, mode="indeterminate", width=200)
        # Progress bar hidden by default
        
        # Show default page
        self._show_recording_page()
    
    def _clear_content(self):
        """Clear all widgets from content frame"""
        for widget in self.content_frame.winfo_children():
            widget.destroy()
    
    def _highlight_nav_button(self, button_text):
        """Highlight the active navigation button"""
        for text, btn in self.nav_buttons.items():
            if text == button_text:
                btn.configure(fg_color=("#3b8ed0", "#1f6aa5"))
            else:
                btn.configure(fg_color="transparent")
    
    # === PAGE METHODS ===
    
    def _show_recording_page(self):
        """Show the recording page"""
        self._clear_content()
        self._highlight_nav_button("🎤 Запись")
        
        # Main container with padding
        container = ctk.CTkFrame(self.content_frame)
        container.grid(row=0, column=0, sticky="nsew", padx=40, pady=30)
        container.grid_columnconfigure(0, weight=1)
        
        # Title
        ctk.CTkLabel(
            container,
            text="Запись встречи",
            font=("Segoe UI", 24, "bold")
        ).grid(row=0, column=0, pady=(0, 30))
        
        # Host selection
        ctk.CTkLabel(
            container,
            text="Ведущий встречи:",
            font=("Segoe UI", 14)
        ).grid(row=1, column=0, sticky="w", pady=(0, 5))
        
        self.host_menu = ctk.CTkOptionMenu(
            container,
            values=[""],
            width=400,
            height=35
        )
        self.host_menu.grid(row=2, column=0, pady=(0, 20))
        self._refresh_host_menu()
        
        # Topic input
        ctk.CTkLabel(
            container,
            text="Тема встречи:",
            font=("Segoe UI", 14)
        ).grid(row=3, column=0, sticky="w", pady=(0, 5))
        
        self.topic_entry = ctk.CTkEntry(
            container,
            width=400,
            height=35,
            placeholder_text="Введите тему..."
        )
        self.topic_entry.grid(row=4, column=0, pady=(0, 30))
        
        # Timer
        self.timer_label = ctk.CTkLabel(
            container,
            text="00:00",
            font=("Segoe UI", 48, "bold")
        )
        self.timer_label.grid(row=5, column=0, pady=20)
        
        # Record button
        self.record_button = ctk.CTkButton(
            container,
            text="⏺ Начать запись",
            width=300,
            height=60,
            font=("Segoe UI", 16, "bold"),
            fg_color="green",
            hover_color="#006400",
            command=self._toggle_recording
        )
        self.record_button.grid(row=6, column=0, pady=20)
        
        # Report button
        self.report_button_live = ctk.CTkButton(
            container,
            text="📄 Сформировать отчёт",
            width=300,
            height=50,
            font=("Segoe UI", 14),
            state="disabled",
            command=self._generate_report
        )
        self.report_button_live.grid(row=7, column=0, pady=10)
    
    def _show_analysis_page(self):
        """Show the analysis page"""
        self._clear_content()
        self._highlight_nav_button("📊 Анализ")
        
        container = ctk.CTkFrame(self.content_frame)
        container.grid(row=0, column=0, sticky="nsew", padx=40, pady=30)
        container.grid_columnconfigure(0, weight=1)
        container.grid_rowconfigure(2, weight=1)
        
        # Title
        ctk.CTkLabel(
            container,
            text="Анализ записи",
            font=("Segoe UI", 24, "bold")
        ).grid(row=0, column=0, pady=(0, 30))
        
        # File selection
        ctk.CTkButton(
            container,
            text="📂 Выбрать аудиофайл",
            width=300,
            height=50,
            font=("Segoe UI", 14),
            command=self._pick_file
        ).grid(row=1, column=0, pady=10)
        
        self.selected_file_label = ctk.CTkLabel(
            container,
            text="",
            font=("Segoe UI", 12),
            wraplength=600
        )
        self.selected_file_label.grid(row=2, column=0, pady=10)
        
        # Report button
        self.report_button_file = ctk.CTkButton(
            container,
            text="📄 Сформировать отчёт",
            width=300,
            height=50,
            font=("Segoe UI", 14),
            state="disabled",
            command=self._generate_report
        )
        self.report_button_file.grid(row=3, column=0, pady=20)
        
        # Transcript area (scrollable)
        ctk.CTkLabel(
            container,
            text="Стенограмма:",
            font=("Segoe UI", 14, "bold")
        ).grid(row=4, column=0, sticky="w", pady=(20, 5))
        
        self.log_box = ctk.CTkTextbox(
            container,
            height=300,
            wrap="word",
            font=("Consolas", 11)
        )
        self.log_box.grid(row=5, column=0, sticky="nsew", pady=5)
    
    def _show_voices_page(self):
        """Show the voices management page"""
        self._clear_content()
        self._highlight_nav_button("🗣️ Голоса")
        self.current_page = "voices"  # Track current page
        
        container = ctk.CTkScrollableFrame(self.content_frame)
        container.grid(row=0, column=0, sticky="nsew", padx=40, pady=30)
        container.grid_columnconfigure(0, weight=1)
        
        # Title
        ctk.CTkLabel(
            container,
            text="База голосов",
            font=("Segoe UI", 24, "bold")
        ).grid(row=0, column=0, pady=(0, 20), sticky="w")
        
        # Buttons
        btn_frame = ctk.CTkFrame(container, fg_color="transparent")
        btn_frame.grid(row=1, column=0, sticky="w", pady=10)
        
        ctk.CTkButton(
            btn_frame,
            text="➕ Добавить голос",
            width=150,
            height=35,
            command=lambda: AddSpeakerDialog(
                self,
                self.voice_db,
                self.recorder,
                self.ai,
                self._log,
                lambda: self._show_voices_page()  # Refresh on change
            )
        ).pack(side="left", padx=5)
        
        ctk.CTkButton(
            btn_frame,
            text="👥 Познакомиться со всеми",
            width=200,
            height=35,
            fg_color="#4c6ef5",
            hover_color="#3b5bdb",
            command=lambda: threading.Thread(target=self._train_all_speakers, daemon=True).start()
        ).pack(side="left", padx=5)
        
        # Speakers list with status
        row = 2
        for name, data in sorted(self.voice_db.items()):
            # Determine status
            if isinstance(data, dict):
                trained = data.get("trained", True)
                sample_count = data.get("sample_count", 0)
            else:
                # Old format (migration should have handled this)
                trained = True
                sample_count = self._get_speaker_sample_count(name)
            
            # Status icon and text
            if trained:
                icon = "✅"
                status_text = f"обучен, {sample_count} семпл(ов)"
                status_color = "#51cf66"
            elif sample_count > 0:
                icon = "⚠️"
                status_text = f"не обучен, {sample_count} семпл(ов)"
                status_color = "#ffd43b"
            else:
                icon = "❌"
                status_text = "не обучен, 0 семплов"
                status_color = "#ff6b6b"
            
            # Speaker row frame
            speaker_frame = ctk.CTkFrame(container)
            speaker_frame.grid(row=row, column=0, sticky="ew", pady=5, padx=10)
            speaker_frame.grid_columnconfigure(1, weight=1)
            
            # Icon + Name
            ctk.CTkLabel(
                speaker_frame,
                text=f"{icon} {name}",
                font=("Segoe UI", 14, "bold")
            ).grid(row=0, column=0, sticky="w", padx=10, pady=10)
            
            # Status
            ctk.CTkLabel(
                speaker_frame,
                text=status_text,
                font=("Segoe UI", 11),
                text_color=status_color
            ).grid(row=0, column=1, sticky="w", padx=10)
            
            # Quick train button (only if untrained but has samples)
            if not trained and sample_count > 0:
                ctk.CTkButton(
                    speaker_frame,
                    text="⚡ Обучить",
                    width=100,
                    height=28,
                    fg_color="#4c6ef5",
                    hover_color="#3b5bdb",
                    command=lambda n=name: threading.Thread(
                        target=self._quick_train_speaker, 
                        args=(n,), 
                        daemon=True
                    ).start()
                ).grid(row=0, column=2, padx=5)
            
            # Re-train button (if trained) - opens dialog to add new sample
            if trained:
                ctk.CTkButton(
                    speaker_frame,
                    text="➕ Добавить семпл",
                    width=130,
                    height=28,
                    fg_color="#868e96",
                    hover_color="#6c757d",
                    command=lambda n=name: AddSpeakerDialog(
                        self,
                        self.voice_db,
                        self.recorder,
                        self.ai,
                        self._log,
                        lambda: self._retrain_after_sample(n),
                        existing_name=n
                    )
                ).grid(row=0, column=2, padx=5)
            
            # Delete button
            ctk.CTkButton(
                speaker_frame,
                text="🗑️",
                width=30,
                height=28,
                fg_color="#c92a2a",
                hover_color="#a61e1e",
                command=lambda n=name: self._remove_speaker_by_name(n)
            ).grid(row=0, column=3, padx=5)
            
            row += 1
    
    def _show_settings_page(self):
        """Show settings page inline"""
        self._clear_content()
        self._highlight_nav_button("⚙️ Настройки")
        
        # Main scrollable container to handle overflow
        container = ctk.CTkScrollableFrame(self.content_frame)
        container.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        container.grid_columnconfigure(0, weight=1)
        
        # Title
        ctk.CTkLabel(
            container,
            text="Настройки",
            font=("Segoe UI", 24, "bold")
        ).grid(row=0, column=0, pady=(10, 20), sticky="w", padx=20)
        
        # Tabs
        tabs = ctk.CTkTabview(container)
        tabs.grid(row=1, column=0, sticky="nsew", padx=10, pady=(0, 10))
        
        tab_sys = tabs.add("Система")
        tab_audio = tabs.add("Аудио")
        tab_stt = tabs.add("Распознавание")
        tab_llm = tabs.add("Нейросеть")
        tab_export = tabs.add("Экспорт")
        
        # --- СИСТЕМА ---
        self._build_settings_entry(tab_sys, "HF Token:", "hf_token", password=True)
        self._build_settings_entry(tab_sys, "Ключевые слова (через запятую):", "keywords")
        
        # Theme selector
        ctk.CTkLabel(
            tab_sys, text="Тема интерфейса:", font=("Segoe UI", 12, "bold")
        ).pack(anchor="w", padx=20, pady=(15, 5))
        self.settings_theme_var = ctk.StringVar(value=ctk.get_appearance_mode())
        ctk.CTkSegmentedButton(
            tab_sys,
            values=["Dark", "Light", "System"],
            variable=self.settings_theme_var,
            command=self._change_theme
        ).pack(fill="x", padx=20, pady=5)
        
        if not FFmpegInstaller.is_installed():
            ctk.CTkButton(
                tab_sys,
                text="Установить FFmpeg",
                command=lambda: threading.Thread(
                    target=lambda: FFmpegInstaller.install(self._log),
                    daemon=True,
                ).start(),
            ).pack(pady=10, padx=20)
        
        # --- АУДИО ---
        ctk.CTkLabel(
            tab_audio, text="Входной микрофон:", font=("Segoe UI", 12, "bold")
        ).pack(anchor="w", padx=20, pady=(10, 5))
        self.settings_device_var = ctk.StringVar(value=config.get("input_device"))
        devices = AudioHelper.get_devices() or ["Default"]
        ctk.CTkOptionMenu(
            tab_audio, variable=self.settings_device_var, values=devices
        ).pack(fill="x", padx=20, pady=5)
        
        ctk.CTkLabel(
            tab_audio, text="Формат записи:", font=("Segoe UI", 12, "bold")
        ).pack(anchor="w", padx=20, pady=(15, 5))
        self.settings_rec_format_var = ctk.StringVar(value=config.get("rec_format"))
        ctk.CTkSegmentedButton(
            tab_audio,
            values=["wav", "mp3"],
            variable=self.settings_rec_format_var,
        ).pack(fill="x", padx=20, pady=5)
        
        # --- РАСПОЗНАВАНИЕ (STT) ---
        ctk.CTkLabel(
            tab_stt, text="Режим обработки:", font=("Segoe UI", 12, "bold")
        ).pack(anchor="w", padx=20, pady=(10, 5))
        self.settings_processing_mode_var = ctk.StringVar(
            value=config.get("processing_mode")
        )
        ctk.CTkSegmentedButton(
            tab_stt,
            values=["cloud", "local"],
            variable=self.settings_processing_mode_var,
            command=self._settings_switch_stt_mode,
        ).pack(fill="x", padx=20, pady=5)
        
        # Cloud-specific settings frame
        self.settings_frame_cloud_stt = ctk.CTkFrame(tab_stt)
        
        self.settings_cloud_use_mp3 = ctk.BooleanVar(
            value=config.get("cloud_use_mp3")
        )
        ctk.CTkCheckBox(
            self.settings_frame_cloud_stt,
            text="Отправлять в Deepgram в виде MP3",
            variable=self.settings_cloud_use_mp3,
        ).pack(anchor="w", padx=20, pady=5)
        
        self._build_settings_entry(self.settings_frame_cloud_stt, "Ключ Deepgram:", "deepgram_key", password=True)
        
        ctk.CTkButton(
            tab_stt,
            text="Проверить модели (Integrity Check)",
            command=self._settings_check_integrity
        ).pack(pady=10, padx=20)
        
        ctk.CTkLabel(
            tab_stt, text="Локальная модель Whisper:", font=("Segoe UI", 12, "bold")
        ).pack(anchor="w", padx=20, pady=(15, 5))
        self.settings_local_size_var = ctk.StringVar(
            value=config.get("local_model_size")
        )
        ctk.CTkOptionMenu(
            tab_stt,
            variable=self.settings_local_size_var,
            values=["tiny", "base", "small", "medium", "large-v3"],
        ).pack(fill="x", padx=20, pady=5)
        
        # Batch size slider
        ctk.CTkLabel(
            tab_stt, text="Размер пакета (batch size):", font=("Segoe UI", 12, "bold")
        ).pack(anchor="w", padx=20, pady=(15, 5))
        
        self.settings_batch_size_var = ctk.IntVar(value=config.get("batch_size", 8))
        
        batch_frame = ctk.CTkFrame(tab_stt)
        batch_frame.pack(fill="x", padx=20, pady=5)
        
        self.settings_batch_size_label = ctk.CTkLabel(
            batch_frame, text=f"Значение: {self.settings_batch_size_var.get()}"
        )
        self.settings_batch_size_label.pack(side="left", padx=10)
        
        self.settings_batch_size_slider = ctk.CTkSlider(
            batch_frame,
            from_=1,
            to=32,
            number_of_steps=31,
            variable=self.settings_batch_size_var,
            command=self._settings_update_batch_label
        )
        self.settings_batch_size_slider.pack(side="left", fill="x", expand=True, padx=10)
        
        # Show/hide cloud settings based on mode
        self._settings_switch_stt_mode(self.settings_processing_mode_var.get())
        
        # --- НЕЙРОСЕТЬ (LLM) ---
        ctk.CTkLabel(
            tab_llm, text="Провайдер LLM:", font=("Segoe UI", 12, "bold")
        ).pack(anchor="w", padx=20, pady=(10, 5))
        self.settings_llm_provider_var = ctk.StringVar(
            value=config.get("llm_provider")
        )
        ctk.CTkSegmentedButton(
            tab_llm,
            values=["openrouter", "local"],
            variable=self.settings_llm_provider_var,
            command=self._settings_switch_llm_provider,
        ).pack(fill="x", padx=20, pady=5)
        
        # OpenRouter frame
        self.settings_frame_or = ctk.CTkFrame(tab_llm)
        self._build_settings_entry(self.settings_frame_or, "OpenRouter API key:", "or_key", password=True)
        
        ctk.CTkLabel(self.settings_frame_or, text="Модель OpenRouter:").pack(anchor="w", padx=10)
        self.settings_or_model_var = ctk.StringVar(value=config.get("or_model"))
        self.settings_or_model_combo = ctk.CTkComboBox(
            self.settings_frame_or,
            variable=self.settings_or_model_var,
            values=[config.get("or_model")] if config.get("or_model") else [],
            width=300
        )
        self.settings_or_model_combo.pack(fill="x", padx=10, pady=5)
        
        ctk.CTkButton(
            self.settings_frame_or,
            text="Загрузить список моделей",
            command=self._settings_load_or_models
        ).pack(pady=5)
        
        # Local LLM frame
        self.settings_frame_local_llm = ctk.CTkFrame(tab_llm)
        
        # Check if Ollama is installed
        ollama_installed = OllamaManager.is_installed()
        
        if not ollama_installed:
            # Show warning if Ollama not installed
            warning_frame = ctk.CTkFrame(self.settings_frame_local_llm, fg_color=("#f39c12", "#e67e22"))
            warning_frame.pack(fill="x", padx=10, pady=10)
            
            ctk.CTkLabel(
                warning_frame,
                text="⚠ Ollama не установлена",
                font=("Segoe UI", 13, "bold"),
                text_color="white"
            ).pack(anchor="w", padx=10, pady=(10, 5))
            
            ctk.CTkLabel(
                warning_frame,
                text="Для использования локальных моделей необходимо установить Ollama.",
                text_color="white",
                wraplength=400
            ).pack(anchor="w", padx=10, pady=(0, 5))
            
            ctk.CTkButton(
                warning_frame,
                text="🌐 Скачать Ollama",
                command=lambda: self._open_url("https://ollama.com/download"),
                fg_color="white",
                text_color="black",
                hover_color="#ecf0f1"
            ).pack(anchor="w", padx=10, pady=(5, 10))
        
        ctk.CTkLabel(
            self.settings_frame_local_llm,
            text="Локальная модель (Ollama):",
        ).pack(anchor="w", padx=10, pady=(5, 0))
        self.settings_local_model_var = ctk.StringVar(
            value=config.get("local_model")
        )
        models = OllamaManager.get_local_models() or ["None"]
        self.settings_local_model_menu = ctk.CTkOptionMenu(
            self.settings_frame_local_llm,
            variable=self.settings_local_model_var,
            values=models,
        )
        self.settings_local_model_menu.pack(fill="x", padx=10, pady=5)
        
        # Only show model shop if Ollama is installed
        if ollama_installed:
            # Collapsible Ollama Model Shop
            self.settings_shop_expanded = ctk.BooleanVar(value=False)
            
            shop_header = ctk.CTkFrame(self.settings_frame_local_llm)
            shop_header.pack(fill="x", padx=10, pady=(10, 0))
            
            self.settings_shop_toggle_btn = ctk.CTkButton(
                shop_header,
                text="▶ Магазин моделей",
                command=self._settings_toggle_shop,
                width=200,
                fg_color="transparent",
                hover_color=("gray70", "gray30"),
                anchor="w"
            )
            self.settings_shop_toggle_btn.pack(side="left", fill="x", expand=True)
            
            # Shop content (initially hidden)
            self.settings_shop_frame = ctk.CTkScrollableFrame(
                self.settings_frame_local_llm, 
                height=300,
                label_text="Доступные модели Ollama"
            )
            # Don't pack yet, will be packed/unpacked on toggle
            
            # Build shop content
            self._build_ollama_shop_content()
        
        self._settings_switch_llm_provider(self.settings_llm_provider_var.get())
        
        # --- Библиотека промптов ---
        frame_prompts = ctk.CTkFrame(tab_llm)
        frame_prompts.pack(fill="x", padx=10, pady=10)
        
        ctk.CTkLabel(
            frame_prompts, text="Сценарий отчёта:", font=("Segoe UI", 12, "bold")
        ).pack(anchor="w")
        
        self.settings_prompts_map = BUILTIN_PROMPTS.copy()
        self.settings_prompts_map.update(config.get("custom_prompts") or {})
        
        self.settings_current_prompt_name_var = ctk.StringVar(
            value=config.get("current_prompt_name")
        )
        self.settings_prompts_menu = ctk.CTkOptionMenu(
            frame_prompts,
            variable=self.settings_current_prompt_name_var,
            values=list(self.settings_prompts_map.keys()),
            command=self._settings_select_prompt,
            width=320,
        )
        self.settings_prompts_menu.pack(pady=5)
        
        self.settings_prompt_text = ctk.CTkTextbox(tab_llm, height=150)
        self.settings_prompt_text.pack(fill="x", padx=20, pady=5)
        self.settings_prompt_text.insert("0.0", config.get("system_prompt") or "")
        
        frame_prompt_buttons = ctk.CTkFrame(tab_llm)
        frame_prompt_buttons.pack(pady=5)
        self.settings_prompt_name_entry = ctk.CTkEntry(
            frame_prompt_buttons, width=180, placeholder_text="Имя нового сценария..."
        )
        self.settings_prompt_name_entry.pack(side="left", padx=5)
        ctk.CTkButton(
            frame_prompt_buttons,
            text="Сохранить как новый",
            command=self._settings_save_prompt,
        ).pack(side="left", padx=5)
        ctk.CTkButton(
            frame_prompt_buttons,
            text="Удалить сценарий",
            fg_color="red",
            command=self._settings_delete_prompt,
        ).pack(side="left", padx=5)
        
        # --- ЭКСПОРТ ---
        self.settings_save_txt_var = ctk.BooleanVar(value=config.get("save_txt"))
        self.settings_save_docx_var = ctk.BooleanVar(value=config.get("save_docx"))
        ctk.CTkCheckBox(
            tab_export, text="Сохранять стенограмму (.txt)", variable=self.settings_save_txt_var
        ).pack(anchor="w", padx=20, pady=5)
        ctk.CTkCheckBox(
            tab_export, text="Сохранять отчёт (.docx)", variable=self.settings_save_docx_var
        ).pack(anchor="w", padx=20, pady=5)
        
        if GDRIVE_AVAILABLE:
            self.settings_use_gdrive_var = ctk.BooleanVar(
                value=config.get("use_gdrive")
            )
            gdrive_cb = ctk.CTkCheckBox(
                tab_export,
                text="Загружать файлы на Google Drive",
                variable=self.settings_use_gdrive_var,
                command=self._settings_on_gdrive_toggle,
            )
            gdrive_cb.pack(anchor="w", padx=20, pady=5)
        
        # Save button at bottom
        ctk.CTkButton(
            container, 
            text="💾 Сохранить настройки", 
            command=self._settings_save,
            width=200,
            height=40,
            font=("Segoe UI", 14, "bold")
        ).grid(row=2, column=0, pady=20)
    
    # === SETTINGS HELPER METHODS ===
    
    def _build_settings_entry(self, parent, label_text, config_key, password=False):
        """Helper to build a settings entry field"""
        frame = ctk.CTkFrame(parent)
        frame.pack(fill="x", padx=10, pady=5)
        ctk.CTkLabel(frame, text=label_text, width=220).pack(
            side="left", padx=(10, 5)
        )
        entry = ctk.CTkEntry(frame, show="*" if password else "")
        entry.pack(side="left", fill="x", expand=True, padx=(0, 10))
        entry.insert(0, config.get(config_key, ""))
        setattr(self, f"settings_entry_{config_key}", entry)
    
    def _open_url(self, url: str):
        """Open URL in default browser"""
        import webbrowser
        webbrowser.open(url)
        self._log(f"Открыта ссылка: {url}")
    
    def _change_theme(self, mode: str):
        """Change the UI theme"""
        ctk.set_appearance_mode(mode)
        self._log(f"Тема изменена на: {mode}")
    
    def _settings_switch_stt_mode(self, mode: str):
        """Show/hide cloud-specific settings based on STT mode"""
        self.settings_frame_cloud_stt.pack_forget()
        if mode == "cloud":
            self.settings_frame_cloud_stt.pack(fill="x", padx=10, pady=5)
    
    def _settings_switch_llm_provider(self, provider: str):
        """Toggle between OpenRouter and local LLM settings"""
        self.settings_frame_or.pack_forget()
        self.settings_frame_local_llm.pack_forget()
        if provider == "openrouter":
            self.settings_frame_or.pack(fill="x", padx=10, pady=5)
        else:
            self.settings_frame_local_llm.pack(fill="x", padx=10, pady=5)
    
    def _settings_select_prompt(self, name: str):
        """Load selected prompt into text area"""
        text = self.settings_prompts_map.get(name, "")
        self.settings_prompt_text.delete("0.0", "end")
        self.settings_prompt_text.insert("0.0", text)
    
    def _settings_save_prompt(self):
        """Save custom prompt"""
        name = self.settings_prompt_name_entry.get().strip()
        text = self.settings_prompt_text.get("0.0", "end").strip()
        if not name or name in BUILTIN_PROMPTS:
            return
        custom = config.get("custom_prompts") or {}
        custom[name] = text
        config.set("custom_prompts", custom)
        config.save()
        
        self.settings_prompts_map = BUILTIN_PROMPTS.copy()
        self.settings_prompts_map.update(custom)
        self.settings_prompts_menu.configure(values=list(self.settings_prompts_map.keys()))
        self.settings_current_prompt_name_var.set(name)
        self._log(f"Сценарий '{name}' сохранён")
    
    def _settings_delete_prompt(self):
        """Delete custom prompt"""
        name = self.settings_current_prompt_name_var.get()
        if name in BUILTIN_PROMPTS:
            messagebox.showwarning("Удаление", "Встроенные сценарии нельзя удалить")
            return
        custom = config.get("custom_prompts") or {}
        if name in custom:
            custom.pop(name)
            config.set("custom_prompts", custom)
            config.save()
        self.settings_prompts_map = BUILTIN_PROMPTS.copy()
        self.settings_prompts_map.update(custom)
        self.settings_prompts_menu.configure(values=list(self.settings_prompts_map.keys()))
        self.settings_current_prompt_name_var.set("Стандартный (умный протокол)")
        self._settings_select_prompt("Стандартный (умный протокол)")
        self._log(f"Сценарий '{name}' удалён")
    
    def _settings_update_batch_label(self, value):
        """Update batch size label when slider changes"""
        self.settings_batch_size_label.configure(text=f"Значение: {int(float(value))}")
    
    def _settings_load_or_models(self):
        """Load OpenRouter models list"""
        models = fetch_openrouter_models()
        if models:
            self.settings_or_model_combo.configure(values=models)
            if not self.settings_or_model_var.get():
                self.settings_or_model_var.set(models[0])
            messagebox.showinfo("OpenRouter", f"Загружено {len(models)} моделей.")
        else:
            messagebox.showerror("OpenRouter", "Не удалось загрузить список моделей.")
    
    def _settings_check_integrity(self):
        """Check model integrity"""
        win = ctk.CTkToplevel(self)
        win.title("Проверка целостности")
        win.geometry("500x400")
        
        log_box = ctk.CTkTextbox(win)
        log_box.pack(fill="both", expand=True, padx=10, pady=10)
        
        def log(msg):
            try:
                if not win.winfo_exists():
                    return
                
                def _update():
                    try:
                        if win.winfo_exists():
                            log_box.insert("end", msg + "\n")
                            log_box.see("end")
                    except Exception:
                        pass
                
                win.after(0, _update)
            except Exception:
                pass
        
        threading.Thread(
            target=lambda: AIProcessor().check_models_integrity(log),
            daemon=True
        ).start()
    
    def _settings_on_gdrive_toggle(self):
        """Handle Google Drive toggle"""
        if self.settings_use_gdrive_var.get():
            def test_auth():
                gdrive = GDriveClient(self._log)
                if not gdrive.auth():
                    self.after(0, lambda: messagebox.showwarning(
                        "Google Drive",
                        "Не удалось авторизоваться.\n"
                        "Убедитесь, что файл credentials.json находится в папке программы."
                    ))
                    self.after(0, lambda: self.settings_use_gdrive_var.set(False))
            threading.Thread(target=test_auth, daemon=True).start()
    
    def _settings_toggle_shop(self):
        """Toggle Ollama model shop visibility"""
        is_expanded = self.settings_shop_expanded.get()
        
        if is_expanded:
            # Collapse
            self.settings_shop_frame.pack_forget()
            self.settings_shop_toggle_btn.configure(text="▶ Магазин моделей")
            self.settings_shop_expanded.set(False)
        else:
            # Expand
            self.settings_shop_frame.pack(fill="both", padx=10, pady=5)
            self.settings_shop_toggle_btn.configure(text="▼ Магазин моделей")
            self.settings_shop_expanded.set(True)
            # Refresh installed models when opening
            self._refresh_ollama_shop()
    
    def _build_ollama_shop_content(self):
        """Build the Ollama model shop content with categories"""
        # Get installed models (full names like 'gpt-oss:20b')
        installed_models_full = OllamaManager.get_local_models() or []
        
        # Helper to check if a model ID is installed (handles versions)
        def is_model_installed(model_id):
            # Check if any installed model starts with this ID
            return any(installed.startswith(model_id) for installed in installed_models_full)
        
        # Progress bar for downloads
        self.settings_shop_progress = ctk.CTkProgressBar(self.settings_shop_frame, mode="determinate")
        self.settings_shop_progress.grid(row=0, column=0, sticky="ew", padx=5, pady=5)
        self.settings_shop_progress.set(0.0)
        self.settings_shop_progress.grid_remove()  # Hidden by default
        
        self.settings_shop_status = ctk.CTkLabel(self.settings_shop_frame, text="")
        self.settings_shop_status.grid(row=1, column=0, sticky="w", padx=5)
        
        # Create tabs for categories
        row_offset = 2
        for group_name, models in OLLAMA_MODELS.items():
            # Category header
            cat_frame = ctk.CTkFrame(self.settings_shop_frame)
            cat_frame.grid(row=row_offset, column=0, sticky="ew", padx=5, pady=(15, 5))
            
            ctk.CTkLabel(
                cat_frame, 
                text=group_name, 
                font=("Segoe UI", 13, "bold")
            ).pack(anchor="w", padx=5, pady=5)
            
            row_offset += 1
            
            # Models in category
            for m in models:
                model_id = m["id"]
                is_installed = is_model_installed(model_id)
                
                model_frame = ctk.CTkFrame(self.settings_shop_frame)
                model_frame.grid(row=row_offset, column=0, sticky="ew", padx=5, pady=2)
                
                # Model name with indicator
                name_text = f"✓ {m['name']}" if is_installed else m["name"]
                label_color = ("#2ecc71", "#27ae60") if is_installed else ("gray10", "gray90")
                
                ctk.CTkLabel(
                    model_frame,
                    text=name_text,
                    text_color=label_color,
                    width=200,
                    anchor="w"
                ).pack(side="left", padx=5)
                
                # Select button
                ctk.CTkButton(
                    model_frame,
                    text="Выбрать",
                    width=80,
                    command=lambda mid=model_id: self._settings_select_ollama_model(mid),
                    state="normal" if is_installed else "disabled"
                ).pack(side="right", padx=2)
                
                # Download button
                download_text = "Обновить" if is_installed else "Скачать"
                ctk.CTkButton(
                    model_frame,
                    text=download_text,
                    width=80,
                    command=lambda mid=model_id: self._settings_download_ollama_model(mid)
                ).pack(side="right", padx=2)
                
                row_offset += 1
    
    def _refresh_ollama_shop(self):
        """Refresh Ollama shop to update installed model indicators"""
        # Clear all model frames (keep progress and status)
        for widget in self.settings_shop_frame.winfo_children():
            if isinstance(widget, ctk.CTkFrame):
                widget.destroy()
        
        # Rebuild content
        self._build_ollama_shop_content()
    
    def _settings_select_ollama_model(self, model_id: str):
        """Select an Ollama model"""
        self.settings_local_model_var.set(model_id)
        self.settings_shop_status.configure(text=f"Выбрана модель: {model_id}")
        self._log(f"Выбрана модель Ollama: {model_id}")
    
    def _settings_download_ollama_model(self, model_id: str):
        """Download an Ollama model"""
        if not OllamaManager.is_installed():
            messagebox.showerror(
                "Ollama",
                "Ollama не установлена или не найдена в PATH.",
            )
            return
        
        self.settings_shop_progress.grid()
        self.settings_shop_progress.set(0.0)
        self.settings_shop_status.configure(text=f"Загрузка модели {model_id}...")
        
        def on_progress(frac, msg):
            self.after(0, lambda: (
                self.settings_shop_progress.set(frac),
                self.settings_shop_status.configure(text=msg)
            ))
        
        def on_done(success, mid):
            def ui():
                if success:
                    self.settings_shop_progress.set(1.0)
                    self.settings_shop_status.configure(text=f"Модель {mid} успешно загружена!")
                    # Auto-select the model
                    self.settings_local_model_var.set(mid)
                    # Refresh shop UI
                    self.after(500, self._refresh_ollama_shop)
                    # Hide progress after a delay
                    self.after(3000, lambda: self.settings_shop_progress.grid_remove())
                else:
                    self.settings_shop_status.configure(text=f"Ошибка загрузки модели {mid}")
                    self.settings_shop_progress.grid_remove()
            
            self.after(0, ui)
        
        threading.Thread(
            target=lambda: OllamaManager.pull_model(model_id, on_progress, on_done),
            daemon=True
        ).start()
    
    def _settings_save(self):
        """Save all settings"""
        # System
        config.set("hf_token", self.settings_entry_hf_token.get())
        config.set("keywords", self.settings_entry_keywords.get())
        
        # Audio
        config.set("input_device", self.settings_device_var.get())
        config.set("rec_format", self.settings_rec_format_var.get())
        
        # STT
        config.set("processing_mode", self.settings_processing_mode_var.get())
        config.set("cloud_use_mp3", self.settings_cloud_use_mp3.get())
        config.set("deepgram_key", self.settings_entry_deepgram_key.get())
        config.set("local_model_size", self.settings_local_size_var.get())
        config.set("batch_size", self.settings_batch_size_var.get())
        
        # LLM
        config.set("llm_provider", self.settings_llm_provider_var.get())
        config.set("or_key", self.settings_entry_or_key.get())
        config.set("or_model", self.settings_or_model_var.get())
        config.set("local_model", self.settings_local_model_var.get())
        
        # Prompts
        config.set("current_prompt_name", self.settings_current_prompt_name_var.get())
        config.set("system_prompt", self.settings_prompt_text.get("0.0", "end").strip())
        
        # Export
        config.set("save_txt", self.settings_save_txt_var.get())
        config.set("save_docx", self.settings_save_docx_var.get())
        if GDRIVE_AVAILABLE:
            config.set("use_gdrive", self.settings_use_gdrive_var.get())
        
        config.save()
        messagebox.showinfo("Настройки", "Настройки сохранены!")
        self._log("Настройки сохранены")
    
    def _show_files_page(self):
        """Show files browser page with tabs, filters, and search"""
        self._clear_content()
        self._highlight_nav_button("📁 Файлы")
        
        # Main container
        container = ctk.CTkFrame(self.content_frame)
        container.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        container.grid_columnconfigure(0, weight=1)
        container.grid_rowconfigure(2, weight=1)  # Table row expands
        
        # === HEADER WITH SEARCH AND CLEAR ===
        header = ctk.CTkFrame(container)
        header.grid(row=0, column=0, sticky="ew", padx=10, pady=10)
        header.grid_columnconfigure(1, weight=1)
        
        # Search buttons
        self.files_search_global_btn = ctk.CTkButton(
            header, text="🔍 Поиск везде",
            command=lambda: self._files_activate_search("global"),
            width=150
        )
        self.files_search_global_btn.grid(row=0, column=0, padx=5)
        
        self.files_search_user_btn = ctk.CTkButton(
            header, text="🔍 У пользователя",
            command=lambda: self._files_activate_search("user"),
            width=150
        )
        self.files_search_user_btn.grid(row=0, column=1, padx=5, sticky="w")
        
        # Search entry (hidden by default)
        self.files_search_var = ctk.StringVar()
        self.files_search_var.trace_add("write", lambda *args: self._files_refresh_list())
        self.files_search_entry = ctk.CTkEntry(
            header, textvariable=self.files_search_var,
            placeholder_text="Введите запрос...",
            width=300
        )
        # Don't grid yet
        
        # Clear filters button (right side)
        self.files_clear_btn = ctk.CTkButton(
            header, text="✖ Очистить фильтры",
            command=self._files_clear_filters,
            fg_color="#e74c3c",
            hover_color="#c0392b",
            width=180
        )
        self.files_clear_btn.grid(row=0, column=3, padx=5, sticky="e")
        header.grid_columnconfigure(2, weight=1)  # Spacer
        
        # === FILTERS ROW ===
        filters = ctk.CTkFrame(container)
        filters.grid(row=1, column=0, sticky="ew", padx=10, pady=5)
        
        # File type filter
        ctk.CTkLabel(filters, text="Тип:").pack(side="left", padx=(10, 5))
        self.files_type_var = ctk.StringVar(value="Все")
        self.files_type_var.trace_add("write", lambda *args: self._files_refresh_list())
        ctk.CTkSegmentedButton(
            filters,
            values=["Все", "Записи", "Транскрипции", "Отчеты"],
            variable=self.files_type_var
        ).pack(side="left", padx=5)
        
        # Sort filter
        ctk.CTkLabel(filters, text="Сортировка:").pack(side="left", padx=(20, 5))
        self.files_sort_var = ctk.StringVar(value="Дата ↓")
        self.files_sort_var.trace_add("write", lambda *args: self._files_refresh_list())
        ctk.CTkOptionMenu(
            filters,
            variable=self.files_sort_var,
            values=["Дата ↓", "Дата ↑", "Имя A-Z", "Имя Z-A"],
            width=120
        ).pack(side="left", padx=5)
        
        # Period filter
        ctk.CTkLabel(filters, text="Период:").pack(side="left", padx=(20, 5))
        self.files_period_var = ctk.StringVar(value="Все")
        self.files_period_var.trace_add("write", lambda *args: self._files_refresh_list())
        ctk.CTkOptionMenu(
            filters,
            variable=self.files_period_var,
            values=["Все", "Сегодня", "Неделя", "Месяц"],
            width=100
        ).pack(side="left", padx=5)
        
        # === HOST TABS (Scrollable) ===
        # Will be filled by _files_build_host_tabs()
        self.files_tabs_frame = ctk.CTkFrame(container)
        self.files_tabs_frame.grid(row=2, column=0, sticky="ew", padx=10, pady=(10, 5))
        
        # === FILES TABLE ===
        self.files_table_frame = ctk.CTkScrollableFrame(container, height=400)
        self.files_table_frame.grid(row=3, column=0, sticky="nsew", padx=10, pady=5)
        self.files_table_frame.grid_columnconfigure(0, weight=1)
        
        # Initialize state variables
        self.files_current_host = "Все файлы"
        self.files_search_mode = None  # None, "global", "user"
        self.files_all_files = []
        
        # Scan and build
        self._files_scan_all()
        self._files_build_host_tabs()
        self._files_refresh_list()
    
    # === FILES BROWSER HELPER METHODS ===
    
    def _files_scan_all(self):
        """Scan all files in Meeting_Records directory"""
        self.files_all_files = []
        
        if not os.path.exists(RECORDS_DIR):
            return
        
        # Iterate through host directories
        for host_name in os.listdir(RECORDS_DIR):
            host_path = os.path.join(RECORDS_DIR, host_name)
            if not os.path.isdir(host_path):
                continue
            
            # Scan each subdirectory
            for subdir_name, file_type in [
                ("Записи", "recording"),
                ("Транскрипции", "transcript"),
                ("Отчеты", "report")
            ]:
                subdir_path = os.path.join(host_path, subdir_name)
                if not os.path.exists(subdir_path):
                    continue
                
                for filename in os.listdir(subdir_path):
                    filepath = os.path.join(subdir_path, filename)
                    if not os.path.isfile(filepath):
                        continue
                    
                    # Get file stats
                    try:
                        stats = os.stat(filepath)
                        file_date = datetime.fromtimestamp(stats.st_mtime)
                        file_size = stats.st_size
                        
                        self.files_all_files.append({
                            "path": filepath,
                            "name": filename,
                            "host": host_name,
                            "type": file_type,
                            "date": file_date,
                            "size": file_size
                        })
                    except Exception:
                        continue
    
    def _files_build_host_tabs(self):
        """Build scrollable host tabs"""
        # Clear existing tabs
        for widget in self.files_tabs_frame.winfo_children():
            widget.destroy()
        
        # Get unique hosts
        hosts = sorted(set(f["host"] for f in self.files_all_files))
        
        # Create scrollable frame for tabs
        tabs_scroll = ctk.CTkScrollableFrame(
            self.files_tabs_frame,
            orientation="horizontal",
            height=50
        )
        tabs_scroll.pack(fill="x", expand=True)
        
        # Apply current filters for counting
        def get_filtered_count(files_list):
            """Apply current filters to get accurate count"""
            filtered = files_list.copy()
            
            # Filter by type
            file_type = self.files_type_var.get()
            type_map = {
                "Записи": "recording",
                "Транскрипции": "transcript",
                "Отчеты": "report"
            }
            if file_type in type_map:
                filtered = [f for f in filtered if f["type"] == type_map[file_type]]
            
            # Filter by period
            period = self.files_period_var.get()
            if period != "Все":
                now = datetime.now()
                if period == "Сегодня":
                    cutoff = now.replace(hour=0, minute=0, second=0, microsecond=0)
                elif period == "Неделя":
                    cutoff = now - timedelta(days=7)
                elif period == "Месяц":
                    cutoff = now - timedelta(days=30)
                
                filtered = [f for f in filtered if f["date"] >= cutoff]
            
            # Filter by search
            search_query = self.files_search_var.get().lower()
            if search_query:
                filtered = [f for f in filtered if search_query in f["name"].lower()]
            
            return len(filtered)
        
        # "All files" tab
        all_count = get_filtered_count(self.files_all_files)
        btn_all = ctk.CTkButton(
            tabs_scroll,
            text=f"Все файлы ({all_count})",
            command=lambda: self._files_select_host("Все файлы"),
            fg_color="#3498db" if self.files_current_host == "Все файлы" else "transparent",
            width=150
        )
        btn_all.pack(side="left", padx=2)
        
        # Host tabs
        for host in hosts:
            host_files = [f for f in self.files_all_files if f["host"] == host]
            count = get_filtered_count(host_files)
            
            btn = ctk.CTkButton(
                tabs_scroll,
                text=f"{host} ({count})",
                command=lambda h=host: self._files_select_host(h),
                fg_color="#3498db" if self.files_current_host == host else "transparent",
                width=150
            )
            btn.pack(side="left", padx=2)
    
    def _files_select_host(self, host_name):
        """Select a host tab"""
        self.files_current_host = host_name
        self._files_build_host_tabs()  # Rebuild to update highlighting
        self._files_refresh_list()
    
    def _files_activate_search(self, mode):
        """Activate search mode (global or user)"""
        self.files_search_mode = mode
        
        # Show search entry
        if mode == "global":
            self.files_search_entry.grid(row=0, column=2, padx=5, sticky="w")
            self.files_search_entry.focus()
        elif mode == "user":
            if self.files_current_host == "Все файлы":
                messagebox.showinfo(
                    "Поиск",
                    "Сначала выберите пользователя (хоста) из табов"
                )
                return
            self.files_search_entry.grid(row=0, column=2, padx=5, sticky="w")
            self.files_search_entry.focus()
    
    def _files_clear_filters(self):
        """Clear all filters and search"""
        self.files_type_var.set("Все")
        self.files_sort_var.set("Дата ↓")
        self.files_period_var.set("Все")
        self.files_search_var.set("")
        self.files_search_mode = None
        self.files_search_entry.grid_forget()
        self.files_current_host = "Все файлы"
        self._files_build_host_tabs()
        self._files_refresh_list()
    
    def _files_refresh_list(self):
        """Refresh the files list with current filters"""
        # Rebuild tabs to update counters with current filters
        self._files_build_host_tabs()
        
        # Clear table
        for widget in self.files_table_frame.winfo_children():
            widget.destroy()
        
        # Apply filters
        filtered_files = self.files_all_files.copy()
        
        # Filter by host
        if self.files_current_host != "Все файлы":
            filtered_files = [f for f in filtered_files if f["host"] == self.files_current_host]
        
        # Filter by type
        file_type = self.files_type_var.get()
        type_map = {
            "Записи": "recording",
            "Транскрипции": "transcript",
            "Отчеты": "report"
        }
        if file_type in type_map:
            filtered_files = [f for f in filtered_files if f["type"] == type_map[file_type]]
        
        # Filter by period
        period = self.files_period_var.get()
        if period != "Все":
            now = datetime.now()
            if period == "Сегодня":
                cutoff = now.replace(hour=0, minute=0, second=0, microsecond=0)
            elif period == "Неделя":
                cutoff = now - timedelta(days=7)
            elif period == "Месяц":
                cutoff = now - timedelta(days=30)
            
            filtered_files = [f for f in filtered_files if f["date"] >= cutoff]
        
        # Filter by search
        search_query = self.files_search_var.get().lower()
        if search_query:
            filtered_files = [f for f in filtered_files if search_query in f["name"].lower()]
        
        # Sort
        sort_mode = self.files_sort_var.get()
        if sort_mode == "Дата ↓":
            filtered_files.sort(key=lambda f: f["date"], reverse=True)
        elif sort_mode == "Дата ↑":
            filtered_files.sort(key=lambda f: f["date"])
        elif sort_mode == "Имя A-Z":
            filtered_files.sort(key=lambda f: f["name"].lower())
        elif sort_mode == "Имя Z-A":
            filtered_files.sort(key=lambda f: f["name"].lower(), reverse=True)
        
        # Display files
        if not filtered_files:
            ctk.CTkLabel(
                self.files_table_frame,
                text="Файлы не найдены",
                font=("Segoe UI", 14),
                text_color="gray"
            ).pack(pady=50)
            return
        
        # Table header
        header = ctk.CTkFrame(self.files_table_frame)
        header.pack(fill="x", padx=5, pady=(5, 10))
        header.grid_columnconfigure(1, weight=1)  # Name column expands
        
        ctk.CTkLabel(header, text="Тип", width=60, font=("Segoe UI", 11, "bold")).grid(row=0, column=0, padx=5)
        ctk.CTkLabel(header, text="Имя файла", width=300, font=("Segoe UI", 11, "bold"), anchor="w").grid(row=0, column=1, padx=5, sticky="w")
        ctk.CTkLabel(header, text="Хост", width=150, font=("Segoe UI", 11, "bold")).grid(row=0, column=2, padx=5)
        ctk.CTkLabel(header, text="Дата", width=120, font=("Segoe UI", 11, "bold")).grid(row=0, column=3, padx=5)
        ctk.CTkLabel(header, text="Размер", width=80, font=("Segoe UI", 11, "bold")).grid(row=0, column=4, padx=5)
        ctk.CTkLabel(header, text="Действия", width=220, font=("Segoe UI", 11, "bold")).grid(row=0, column=5, padx=5)
        
        # File rows
        for file_data in filtered_files:
            row = ctk.CTkFrame(self.files_table_frame)
            row.pack(fill="x", padx=5, pady=2)
            row.grid_columnconfigure(1, weight=1)
            
            # Icon
            icon_map = {
                "recording": "🎙️",
                "transcript": "📄",
                "report": "📊"
            }
            icon = icon_map.get(file_data["type"], "📁")
            ctk.CTkLabel(row, text=icon, width=60).grid(row=0, column=0, padx=5)
            
            # Name
            name_label = ctk.CTkLabel(row, text=file_data["name"], width=300, anchor="w")
            name_label.grid(row=0, column=1, padx=5, sticky="w")
            
            # Host
            ctk.CTkLabel(row, text=file_data["host"], width=150).grid(row=0, column=2, padx=5)
            
            # Date
            date_str = file_data["date"].strftime("%d.%m.%Y %H:%M")
            ctk.CTkLabel(row, text=date_str, width=120).grid(row=0, column=3, padx=5)
            
            # Size
            size_mb = file_data["size"] / (1024 * 1024)
            size_str = f"{size_mb:.1f} MB" if size_mb >= 1 else f"{file_data['size'] // 1024} KB"
            ctk.CTkLabel(row, text=size_str, width=80).grid(row=0, column=4, padx=5)
            
            # Actions
            actions = ctk.CTkFrame(row)
            actions.grid(row=0, column=5, padx=5)
            
            ctk.CTkButton(
                actions, text="Открыть", width=60,
                command=lambda p=file_data["path"]: self._files_open_file(p)
            ).pack(side="left", padx=2)
            
            ctk.CTkButton(
                actions, text="Папка", width=60,
                command=lambda p=file_data["path"]: self._files_show_in_folder(p)
            ).pack(side="left", padx=2)
            
            ctk.CTkButton(
                actions, text="Удалить", width=70,
                fg_color="#e74c3c", hover_color="#c0392b",
                command=lambda p=file_data["path"]: self._files_delete_file(p)
            ).pack(side="left", padx=2)
    
    def _files_open_file(self, filepath):
        """Open file in default application"""
        try:
            os.startfile(filepath)
            self._log(f"Открыт файл: {os.path.basename(filepath)}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть файл:\n{e}")
    
    def _files_show_in_folder(self, filepath):
        """Show file in Windows Explorer"""
        try:
            subprocess.run(["explorer", "/select,", filepath])
            self._log(f"Показан в проводнике: {os.path.basename(filepath)}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть проводник:\n{e}")
    
    def _files_delete_file(self, filepath):
        """Delete file with confirmation"""
        filename = os.path.basename(filepath)
        confirm = messagebox.askyesno(
            "Подтверждение",
            f"Удалить файл?\n\n{filename}"
        )
        
        if confirm:
            try:
                os.remove(filepath)
                self._log(f"Удалён файл: {filename}")
                # Refresh
                self._files_scan_all()
                self._files_build_host_tabs()
                self._files_refresh_list()
                messagebox.showinfo("Успех", f"Файл удалён:\n{filename}")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось удалить файл:\n{e}")

    def _update_speakers_box(self):
        self.speakers_box.configure(state="normal")
        self.speakers_box.delete("0.0", "end")
        for name in self.voice_db.keys():
            self.speakers_box.insert("end", f"{name}\n")
        self.speakers_box.configure(state="disabled")

    def _on_speakers_changed(self):
        self._save_db()
        self._update_speakers_box()
        self._refresh_host_menu()

    def _refresh_host_menu(self):
        names = list(self.voice_db.keys()) or ["Добавьте участника"]
        if hasattr(self, "host_menu"):
            self.host_menu.configure(values=names)
            self.host_menu.set(names[0])

    # --- управление записью ---

    def _toggle_recording(self):
        if not self.is_recording:
            host = self.host_menu.get()
            if not host or "Добавьте" in host:
                self._log("Нужно выбрать ведущего встречи.")
                return

            topic = self.topic_entry.get().strip() or "Встреча"
            safe_topic = sanitize_filename(topic)
            date_prefix = datetime.now().strftime("%Y-%m-%d")
            
            # Get host-specific directories
            rec_dir, _, _ = self._get_host_dirs(host)
            
            filename = f"{date_prefix}_{safe_topic}_{sanitize_filename(host)}.{config.get('rec_format')}"
            full_path = os.path.join(rec_dir, filename)

            self.recorder.start(full_path)
            self.is_recording = True
            self.record_start_time = time.time()
            self.record_button.configure(
                text="Остановить запись", fg_color="red"
            )
            self._update_timer()
        else:
            self.recorder.stop()
            self.is_recording = False
            self.record_button.configure(
                text="Начать запись", fg_color="green"
            )
            # анализ только что записанного файла
            full_path = self.recorder.final_target
            self._start_analysis(full_path)

    def _update_timer(self):
        if self.is_recording:
            elapsed = int(time.time() - self.record_start_time)
            self.timer_label.configure(
                text=f"{elapsed // 60:02}:{elapsed % 60:02}"
            )
            self.after(1000, self._update_timer)

    # --- обработка файла ---

    def _pick_file(self):
        path = filedialog.askopenfilename()
        if not path:
            return
        self.selected_file_label.configure(text=os.path.basename(path))
        self._start_analysis(path)

    def _start_analysis(self, path: str):
        self._log("Запуск анализа аудио...")
        self.progress.pack(side="right", padx=10)
        self.progress.start()
        threading.Thread(target=self._run_analysis, args=(path,), daemon=True).start()

    def _run_analysis(self, path: str):
        try:
            try:
                logger.info(f"Starting analysis of: {path}")
                segments, unknown = self.ai.analyze(path, self.voice_db, self._log)
                logger.info(f"Analysis completed. Segments: {len(segments)}, Unknown: {len(unknown)}")
            except Exception as e:
                self._log(f"Ошибка анализа: {e}")
                logger.error(f"Analysis failed: {e}", exc_info=True)
                self._stop_progress()
                return

            if unknown:
                self._log("Обнаружены неизвестные голоса, запускаю мастер...")
                wizard = IdentifyWizard(self, unknown, self.voice_db)
                self.wait_window(wizard)
                # сохраняем новые голоса
                for spk_id, name in wizard.result_names.items():
                    if wizard.save_flags.get(spk_id):
                        sample = unknown[spk_id]
                        emb = self.ai.create_embedding(sample["audio"])
                        self.voice_db[name] = emb
                self._save_db()
                self._update_speakers_box()
                self._refresh_host_menu()
                # переименуем метки в сегментах
                for seg in segments:
                    label = seg["label"]
                    if label in wizard.result_names:
                        seg["label"] = wizard.result_names[label]

            # формируем стенограмму
            transcript = "\n".join(
                f"[{s['start']:.1f}] {s['label']}: {s['text']}"
                for s in segments
                if s["text"].strip()
            )
            self.last_transcript_text = transcript
            self.last_basename = os.path.splitext(os.path.basename(path))[0]

            if config.get("save_txt"):
                # Try to extract host from filename (format: Date_Topic_Host)
                # Fallback to current selected host in UI if available
                host_name = "Unknown_Host"
                try:
                    parts = self.last_basename.split("_")
                    if len(parts) >= 3:
                        host_name = parts[-1]
                    elif hasattr(self, "host_menu"):
                        val = self.host_menu.get()
                        if val and "Добавьте" not in val:
                            host_name = val
                except Exception:
                    pass

                _, trans_dir, _ = self._get_host_dirs(host_name)
                
                txt_path = os.path.join(
                    trans_dir, f"{self.last_basename}.txt"
                )
                with open(txt_path, "w", encoding="utf-8") as f:
                    f.write(transcript)
                self._log(f"Сохранена стенограмма: {txt_path}")
                if config.get("use_gdrive"):
                    threading.Thread(
                        target=lambda: self.gdrive.upload(txt_path),
                        daemon=True,
                    ).start()

            # включаем кнопки отчёта
            self.after(
                0,
                lambda: (
                    self.report_button_live.configure(state="normal"),
                    self.report_button_file.configure(state="normal"),
                ),
            )
            self._stop_progress()

        finally:
            # Cleanup temp directory
            if os.path.exists(TEMP_DIR):
                try:
                    shutil.rmtree(TEMP_DIR)
                    logger.info(f"Cleaned up temp directory: {TEMP_DIR}")
                except Exception as e:
                    logger.warning(f"Failed to cleanup temp directory: {e}")

    def _stop_progress(self):
        self.after(
            0,
            lambda: (
                self.progress.stop(),
                self.progress.pack_forget(),
            ),
        )

    # --- генерация отчёта LLM ---

    def _generate_report(self):
        if not self.last_transcript_text:
            self._log("Нет стенограммы для отчёта.")
            return

        self.progress.pack(side="right", padx=10)
        self.progress.start()
        # Show status (assuming status_label exists or using log)
        self._log("Генерация отчёта...")
        
        threading.Thread(target=self._llm_report_thread, daemon=True).start()

    def _llm_report_thread(self):
        report = self.llm.summarize(self.last_transcript_text)
        if config.get("save_docx"):
            # Determine host similar to _run_analysis
            host_name = "Unknown_Host"
            try:
                parts = self.last_basename.split("_")
                if len(parts) >= 3:
                    host_name = parts[-1]
                elif hasattr(self, "host_menu"):
                    val = self.host_menu.get()
                    if val and "Добавьте" not in val:
                        host_name = val
            except Exception:
                pass

            _, _, rep_dir = self._get_host_dirs(host_name)
            
            docx_path = os.path.join(
                rep_dir, f"{self.last_basename}_report.docx"
            )
            DocxGenerator.create_report(report, docx_path)
            self._log(f"Сохранён отчёт: {docx_path}")
            if config.get("use_gdrive"):
                self.gdrive.upload(docx_path)
        self._stop_progress()

    # --- работа с базой спикеров ---

    def _remove_speaker(self):
        if not self.voice_db:
            return
        win = ctk.CTkToplevel(self)
        win.title("Удаление участника")
        win.geometry("300x120")
        ctk.CTkLabel(win, text="Кого удалить?").pack(pady=10)
        var = ctk.StringVar(value=list(self.voice_db.keys())[0])
        ctk.CTkOptionMenu(win, variable=var,
                          values=list(self.voice_db.keys())
                          ).pack(pady=5)

        def do_delete():
            self.voice_db.pop(var.get(), None)
            self._save_db()
            self._update_speakers_box()
            self._refresh_host_menu()
            win.destroy()

        ctk.CTkButton(
            win, text="Удалить", fg_color="red", command=do_delete
        ).pack(pady=10)

    # === SPEAKER TRAINING METHODS ===
    
    def _quick_train_speaker(self, name):
        """Train or re-train a speaker using all their samples"""
        import glob
        
        speaker_dir = os.path.join(SAMPLES_DIR, sanitize_filename(name))
        samples = sorted(glob.glob(os.path.join(speaker_dir, "sample_*.wav")))
        
        if not samples:
            self._log(f"Нет семплов для {name}")
            return
        
        self._log(f"Обучение {name} на основе {len(samples)} семпл(ов)...")
        
        # Create embeddings from all samples and average them
        embeddings = []
        for sample_path in samples:
            try:
                emb = self.ai.create_embedding(sample_path)
                embeddings.append(emb)
            except Exception as e:
                logger.warning(f"Failed to process {sample_path}: {e}")
        
        if not embeddings:
            self._log(f"❌ Не удалось обработать семплы для {name}")
            return
        
        # Average embeddings for better accuracy
        avg_embedding = np.mean(embeddings, axis=0)
        
        # Update voice_db
        if isinstance(self.voice_db.get(name), dict):
            self.voice_db[name]["embedding"] = avg_embedding
            self.voice_db[name]["trained"] = True
            self.voice_db[name]["sample_count"] = len(samples)
        else:
            self.voice_db[name] = {
                "embedding": avg_embedding,
                "trained": True,
                "sample_count": len(samples)
            }
        
        self._save_db()
        # Update UI if on voices page
        if hasattr(self, 'current_page') and self.current_page == "voices":
            self._show_voices_page()
        self._log(f"✅ {name} обучен (использовано {len(embeddings)} семпл(ов))")
    
    def _train_all_speakers(self):
        """Batch train all untrained speakers with samples"""
        untrained = []
        for name, data in self.voice_db.items():
            if isinstance(data, dict):
                is_trained = data.get("trained", True)
                has_samples = data.get("sample_count", 0) > 0
                if not is_trained and has_samples:
                    untrained.append(name)
        
        if not untrained:
            self._log("✅ Все спикеры уже обучены!")
            return
        
        self._log(f"Начинаю обучение {len(untrained)} спикеров...")
        
        for i, name in enumerate(untrained, 1):
            self._log(f"[{i}/{len(untrained)}] {name}...")
            self._quick_train_speaker(name)
        
        self._log(f"🎉 Пакетное обучение завершено! Обучено: {len(untrained)}")
    
    def _remove_speaker_by_name(self, name):
        """Remove a speaker by name with confirmation"""
        import shutil
        
        result = messagebox.askyesno(
            "Удалить спикера",
            f"Удалить {name} из базы?\nВсе семплы также будут удалены."
        )
        if result:
            # Remove from database
            if name in self.voice_db:
                self.voice_db.pop(name)
                self._save_db()
            
            # Remove samples folder
            speaker_dir = os.path.join(SAMPLES_DIR, sanitize_filename(name))
            if os.path.exists(speaker_dir):
                shutil.rmtree(speaker_dir)
            
            self._log(f"Удалён: {name}")
            # Refresh page
            self._show_voices_page()
    
    def _retrain_after_sample(self, name):
        """Called after adding a new sample - retrain the speaker"""
        threading.Thread(
            target=self._quick_train_speaker,
            args=(name,),
            daemon=True
        ).start()
        # Refresh voices page
        self._show_voices_page()


if __name__ == "__main__":
    app = App()
    app.mainloop()
